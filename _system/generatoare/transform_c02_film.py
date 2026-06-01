#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Transforma FILM-Excel-02 din MARCARE in CONTROL. Inlocuire pe index de paragraf."""
import docx, os

SRC = 'c02/FILM-Excel-02-Marcare.docx'
DST = 'c02/FILM-Excel-02-Control.docx'

REPL = {
 5: "Setul de date pare valid. Dar valid nu înseamnă corect.",
 7: "Valorile valide pot fi greșite față de realitate. C02 le confruntă cu nomenclatoare, reguli și calendare, semnalează fiecare neconcordanță cu motivul ei și nu modifică nicio valoare sursă.",
 9: "Nu credem valorile. Le confruntăm cu realitatea.",
 11: "Excel acceptă orice valoare validă. Acum ai aparatul care întreabă dacă e și corectă.",
 13: "Încrederea începe după confruntarea cu realitatea.",
 16: "BOMBA (idee dominantă): Pare valid. E greșit.",
 17: "SUNĂ CUNOSCUT (recunoaștere): Lucrezi cu un set care a trecut toate validările. Și totuși: un oraș scris în trei feluri diferite; o cotă de TVA care nu respectă categoria; o scadență căzută înaintea facturii.",
 18: "GREȘEALA CLASICĂ: Oamenii cred valoarea fiindcă Excel o acceptă. Profesioniștii o confruntă cu realitatea.",
 19: "AHA (vârf): Valid nu înseamnă corect.",
 20: "CINE DEVII (promisiune): Nu mai accepți valoarea pe cuvânt. O confrunți cu realitatea.",
 25: "Tabelul pare valid",
 27: "Aparența validității. Astăzi am descoperit de ce era doar aparență.",
 35: "Construirea semnalizării",
 37: "Nu am modificat valori. Le-am confruntat cu realitatea.",
 40: "Verificarea corespondenței",
 42: "O semnalare scrisă nu e o semnalare confirmată. COUNTIF decide.",
 45: "Ancorare la sursă",
 47: "Un set confruntat azi nu se confruntă singur luna viitoare. Decât dacă formulele rămân ancorate.",
 52: "Valorile păreau valide. De data asta, le-am confruntat pe fiecare cu realitatea.",
 54: "Constructia 02 CONTROL preia setul predat de constructia precedenta. C01 l-a predat ca structura corecta. Valorile insa pot fi valide tehnic si totusi gresite fata de realitate: orase care nu exista in judetul lor, cote de TVA nelegale pentru categorie, scadente inaintea facturii, facturi in zile nelucratoare, CNP-uri care se contrazic intern.",
 55: "Setul are acum coloane de semnalizare: flag, status_validare si motiv_anomalie. Fiecare neconcordanta fata de realitate e semnalata, motivata, verificabila.",
 59: "Interogare + semnalizare. Raporteaza ce vede, construieste formulele de confruntare cerute explicit de noi.",
 60: "EXCEL / FORMULE",
 61: "Furnizeaza functii native de confruntare: COUNTIF, XLOOKUP, NETWORKDAYS. Valorile sursa raman neatinse.",
 62: "Cerem demonstratia, confruntam fiecare valoare cu realitatea, decidem. Nu reparam, doar semnalam cu dovezi.",
 63: "SCENA REALA · CELE 5 FENOMENE",
 64: "01 · ORASE NEALINIATE",
 65: "oras inexistent in judet sau scris diferit, ~50 cazuri",
 66: "02 · TVA GRESIT",
 67: "cota valida ca numar, nelegala pentru categorie, ~35 cazuri",
 68: "03 · SCADENTA INAINTEA FACTURII",
 69: "data_scadentei < data_factura, ~25 cazuri",
 70: "04 · ZI NELUCRATOARE",
 71: "factura in sarbatoare legala sau duminica, ~20 cazuri",
 72: "05 · CNP CONTRADICTORIU",
 73: "CNP care se contrazice cu sex, judet, data nasterii sau cifra de control, ~25 contacte",
 74: "Total: ~130 randuri semnalate in 5 categorii. Valorile sursa raman neatinse (R-V02.14).",
 78: "Deschidem setul predat de C01, recunoastem ca validul poate fi gresit, deschidem nomenclatoarele de referinta.",
 95: "Promptul 1 interogheaza setul: ce valori par valide dar gresite fata de realitate, pe cinci categorii. Nicio modificare.",
 112: "Promptul 2 construieste aparatul de semnalizare cu formule: COUNTIF pe localitati, XLOOKUP pe TVA, comparatie de date, NETWORKDAYS pe calendar. Valorile nu se modifica.",
 129: "COUNTIF pe fiecare flag confrunta semnalarile cu setul si le cuantifica.",
 146: "Tabelul structurat ancoreaza formulele. Orice rand nou primeste aceeasi confruntare cu realitatea.",
 163: "Setul cu anomaliile semnalate predat catre constructia urmatoare cap-coada. Valorile sursa neatinse.",
}

def set_para(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)

d = docx.Document(SRC)
n = 0
for i, p in enumerate(d.paragraphs):
    if i in REPL:
        set_para(p, REPL[i]); n += 1
    elif 'Confirmam martorul:' in p.text:
        # MOMENT DE CONTROL: scoatem ecoul de suma-martor (R-TERM-5)
        set_para(p, p.text.replace('Confirmam martorul:', 'Confirmam controlul:')); n += 1
d.save(DST)

# verificare leftover MARCARE-isme
import re
full = "\n".join(p.text for p in docx.Document(DST).paragraphs)
left = []
for t in ['Marcare','marchează','marca','week-end','duplicat','cod client','SUMIF','AGGREGATE','Power Query','suma martor','martor','câmp obligatoriu','36 anomalii','adevărat','viitor']:
    c = len(re.findall(t, full))
    if c: left.append(f"{t}={c}")
print(f"FILM: {n} paragrafe inlocuite -> {DST}")
print("leftover:", ", ".join(left) if left else "ZERO")
