#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c15.py - FILM-Excel-15-Sintetizare.docx (script video procedural).

Structura = c14 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C15 SINTETIZARE.
Garda C15: a formula mesajul esential (o propozitie). SINTEZA != REZUMAT. ZERO calcul (T3),
zero rearanjare pagina (C14), zero livrare/decizie (C16). Cifre in Excel (R-V02.15).
Corp narativ pe vocea axei SINTEZA (L200): mesaj, fraza, sinteza, headline, so-what.
E5 = REFORMULARE.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = 'c15/FILM-Excel-15-Sintetizare.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    os.makedirs('c15', exist_ok=True)
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C15 SINTETIZARE')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 4 RAPORTARE · Construcția 15 din 20 · a treia din T4')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK->DEMO->CONTROL->REVEAL.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('O pagină impecabilă și niciun mesaj. O deschizi și nu știi ce să reții. Astăzi îi extragem o singură frază: ce contează, dintr-o privire.')
    lbl('MIZA')
    txt('Un director nu citește toată pagina; o deschide și întreabă și ce-i cu asta? Trecem de la o '
        'pagină corectă care arată tot și nu spune nimic, la o singură frază care spune ce contează, '
        'dovedită de pagina de dedesubt. Nu o cifră nouă, nu o cauză nouă, acelea vin gata din analiză, '
        'ci mesajul: ce trebuie reținut.')
    lbl('AHA (INSIGHT)')
    txt('O pagină arată. O sinteză spune.')
    lbl('MANTRA · TRAINITY')
    txt('Nu rezumăm. Sintetizăm.')
    lbl('WOW (PAYOFF FINAL)')
    txt('O pagină întreagă pe care trebuia s-o descifrezi. Acum o singură frază îți spune ce contează, iar pagina o dovedește.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Dintr-o privire, mesajul.')
    lbl('FORMULA')
    txt('Rezumatul scurtează tot. Sinteza spune ce contează.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion. FRAZĂ DE IMPACT = exec-phrase.')
    execs = [
        ('1 · MUTĂ', 'Pagina mută',
         'O pagină coerentă, corectă, completă. O privești trei secunde și tot nu știi ce să reții.'),
        ('2 · REZUMATUL', 'Rezumatul automat',
         'AI scurtează pagina. Rezumatul scurtează tot; sinteza spune ce contează. Nu sunt același lucru.'),
        ('3 · MESAJUL', 'Mesajul esențial',
         'Alegi singura afirmație pe care pagina o dovedește și o formulezi într-o frază, pentru decizia de față.'),
        ('4 · TESTUL', 'Testul și ce-i cu asta?',
         'Fără această frază, decidentul hotărăște la fel? Indicativ, nu decizional. Nimic nou peste pagină.'),
        ('5 · REFORMULAREA', 'Mesajul se adaptează',
         'Datele s-au schimbat în amonte. Reformulezi headline-ul, refaci sinteza, nu recalculezi nimic.'),
        ('6 · PREDAREA', 'O pagină a devenit un mesaj',
         'Headline-ul sus, pagina-dovadă dedesubt. Gata de predat pentru încadrarea deciziei la livrare.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C15')
    txt('Construcția 15 SINTETIZARE preia pagina coerentă de la treapta de compunere. Avem un focar, '
        'o ierarhie, un traseu de citire, totul corect. Dar o pagină arată; nu spune. Mesajul esențial '
        'există în date, dar nu e formulat nicăieri. Acum nu mai rearanjăm nimic; extragem o frază.')
    txt('Pentru prima dată, miza nu mai e cum arată pagina, ci ce spune. Rezumatul scurtează tot. '
        'Sinteza alege singurul mesaj care contează. C15 formulează acel mesaj.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Propune câteva titluri candidate pornind de la raport și critică titlul ales. Scurtează și '
        'verifică repede. Nu alege mesajul; alegerea a ce contează o face omul.')
    lbl('EXCEL · SINTEZA')
    txt('Găzduiește foaia Sinteza, stratul MESAJ: headline-ul formulat de cursant, linia so-what, '
        'legătura la pagina-dovadă și reformularea. Citește din pagină, nu produce pagina; nu calculează.')
    lbl('CONTROL UMAN')
    txt('Alegem singura afirmație pe care pagina o dovedește, o formulăm într-o frază pentru această '
        'decizie și audiență, o trecem prin testul so-what. Nu calculăm și nu rearanjăm: formulăm mesajul.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CELE 5 FENOMENE')
    fenomene = [
        ('01 · REZUMATUL CARE NU SPUNE',
         'Scurtezi tot proporțional și obții o listă fără ierarhie de sens, nimic nu iese în față. Operația: sinteza alege singurul mesaj care contează, nu îl scurtează pe tot.'),
        ('02 · MESAJUL ÎNGROPAT',
         'Pagina arată totul, dar afirmația-cheie nu e scrisă nicăieri, cititorul o caută singur. Operația: o singură frază pe care pagina de dedesubt o dovedește.'),
        ('03 · ȘI CE-I CU ASTA?',
         'Decidentul citește pagina întreagă și tot întreabă care e ideea. Operația: fraza care trece testul fără ea, decidentul hotărăște la fel?'),
        ('04 · MESAJUL CARE NU SE POTRIVEȘTE',
         'Aceeași pagină, dar mesajul e scris pentru altă audiență sau altă decizie. Operația: mesajul pentru această decizie și această audiență, nu unul în general.'),
        ('05 · AFIRMAȚIA NESUSȚINUTĂ',
         'Scrii o frază pe care pagina nu o dovedește, mesaj fără acoperire. Operația: headline ancorat în dovada vizuală deja prezentă pe pagină.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 5 fenomene de sinteză, fiecare cu operația lui de formulare. Pagina e pregătită '
        'să capete o singură frază care contează.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · MUTĂ',
            'obiectiv': 'Preluăm pagina coerentă de la treapta de compunere, o privim trei secunde și constatăm că o pagină corectă poate să nu spună nimic.',
            'stare': 'Atenție. Pagină corectă, completă; o privești și tot nu știi ce să reții.',
            'voce': '„Etapa 1: pagina mută. Avem o pagină coerentă: focar, ierarhie, traseu. Totul se vede. Și totuși nu spune nimic. O pagină arată; nu spune."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C15, foaia Sinteza; arătăm pagina compusă, corectă, fără un mesaj scris.',
            'ai': 'AI-ul confirmă că pagina e corectă și completă. Noi observăm că lipsește o frază, fără să atingem datele.',
            'control': 'Punem regula-test: mesajul e o afirmație în cuvinte, nu o sumă de obiecte corecte. „O pagină arată; nu spune. Îi dăm o frază."',
            'reveal': 'Pe ecran: pagina coerentă, corectă, dar mută.',
        },
        {
            'nume': 'ETAPA 2 · REZUMATUL',
            'obiectiv': 'Lăsăm AI-ul să scurteze pagina și vedem diferența dintre a rezuma tot și a spune ce contează. Promptul 1.',
            'stare': 'Curiozitate metodică. Primim materie primă, nu mesajul.',
            'voce': '„Etapa 2: rezumatul. AI face partea automatizabilă, un rezumat draft. Dar rezumatul scurtează tot. Sinteza spune ce contează. Nu e același lucru."',
            'demo': 'Pe ecran: rulăm Promptul 1; AI-ul dă trei titluri candidate; le punem lângă un rezumat clasic.',
            'ai': 'Promptul 1 propune titluri candidate din ce e deja în raport, fără cifre noi. Noi alegem și ascuțim; AI-ul nu alege mesajul.',
            'control': 'Confirmăm distincția: „Rezumatul scurtează tot. Sinteza spune ce contează. Mesajul nu e ce conține pagina, e ce trebuie reținut din ea."',
            'reveal': 'Pe ecran: un rezumat proporțional alături de candidați de mesaj; diferența e vizibilă.',
        },
        {
            'nume': 'ETAPA 3 · MESAJUL',
            'obiectiv': 'Alegem singura afirmație pe care pagina o dovedește și o formulăm într-o frază, pentru decizia și audiența de față.',
            'stare': 'Concentrare. Aici se naște fraza, partea cea mai umană.',
            'voce': '„Etapa 3: mesajul. Care e singura afirmație pe care pagina o dovedește? O formulăm într-o frază. Mesajul nu se descoperă, vine din analiză; noi îl formulăm."',
            'demo': 'Pe ecran: ne uităm la focarul paginii, scriem headline-ul și linia so-what în foaia Sinteza, cu legătura la dovadă.',
            'ai': 'AI-ul poate reformula fraza, dar judecata mesajului potrivit pentru această decizie e a omului.',
            'control': 'Confirmăm: o singură afirmație, susținută de pagină, pentru această audiență. „Aceeași pagină poate avea mesaje diferite; sinteza alege mesajul potrivit."',
            'reveal': 'Pe ecran: headline-ul de o frază, cu pagina-dovadă dedesubt.',
        },
        {
            'nume': 'ETAPA 4 · TESTUL',
            'obiectiv': 'Trecem mesajul prin trei probe: contează, e indicativ și nu inventează nimic peste pagină. Promptul 2.',
            'stare': 'Rigoare. Mesajul trece sau cade la filtru.',
            'voce': '„Etapa 4: testul. Fără această frază, decidentul hotărăște la fel? Dacă da, nu e mesajul. Indicativ, nu decizional. Nimic nou peste pagină."',
            'demo': 'Pe ecran: rulăm Promptul 2; AI-ul critică titlul; verificăm o singură afirmație, indicativ, fără cifră sau cauză nouă.',
            'ai': 'Promptul 2 verifică mesajul după patru criterii și spune unde cade. Noi luăm decizia finală asupra formei.',
            'control': 'Aplicăm filtrul so-what. „C15 enunță; nu numește decizia, nu recomandă. Nicio cifră nouă, nicio cauză nouă; acelea vin din analiză."',
            'reveal': 'Pe ecran: mesajul care trece testul, alături de o frază-zgomot care cade.',
        },
        {
            'nume': 'ETAPA 5 · REFORMULAREA',
            'obiectiv': 'Când datele se schimbă în amonte, reformulăm headline-ul în cuvinte, fără să recalculăm nimic.',
            'stare': 'Adaptare calmă. Mesajul ține pasul cu pagina.',
            'voce': '„Etapa 5: reformularea. Datele s-au actualizat în amonte. Cifrele s-au schimbat acolo. Reformulăm fraza, refacem sinteza. Nu recalculăm nimic."',
            'demo': 'Pe ecran: pagina reflectă date noi; rescriem headline-ul în foaia Sinteza; cifrele rămân ale modelului.',
            'ai': 'AI-ul poate propune un nou titlu pe pagina actualizată. Noi confirmăm că am reformulat, nu recalculat.',
            'control': 'Confirmăm: „Reformularea e un act în cuvinte, nu un recalcul. Noul mesaj rămâne o singură afirmație susținută, pe criteriul ce contează."',
            'reveal': 'Pe ecran: headline vechi tăiat, headline nou scris, aceleași cifre ale modelului.',
        },
        {
            'nume': 'ETAPA 6 · PREDAREA',
            'obiectiv': 'Punem mesajul în capul paginii, verificăm că fraza și dovada stau împreună și predăm pentru încadrarea deciziei.',
            'stare': 'Claritate. Pagina nu mai e mută; spune o frază.',
            'voce': '„Etapa 6: predarea. Headline-ul sus, dovada dedesubt. Aceeași pagină care doar arăta, acum spune. Nu mai prezentăm tot; sintetizăm fraza care contează."',
            'demo': 'Pe ecran: punem headline-ul în capul paginii; predarea paginii cu mesaj către treapta de livrare.',
            'ai': 'AI-ul rezumă mesajul final. Noi confirmăm că am formulat mesajul, nu am numit decizia.',
            'control': 'Confirmăm: o frază, dovedită de pagină, pentru această decizie. „C15 spune ce contează. Încadrarea pentru decizie vine la treapta de livrare."',
            'reveal': 'Pe ecran: pagina cu headline, predată; o pagină a devenit un mesaj care contează.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · O PAGINĂ A DEVENIT UN MESAJ')
    txt('C15 extrage și formulează mesajul esential al paginii coerente primite de la treapta de '
        'compunere. O pagină care doar arăta capătă o singură frază pe care pagina o dovedește, '
        'trecută prin testul so-what și reformulată când pagina se schimbă. C15 spune ce contează; '
        'încadrarea pentru decizie începe la treapta de livrare. Predăm un raport care spune, cu '
        'valorile sursă neatinse și suma conservată cap-coadă.')
    p = doc.add_paragraph()
    r = p.add_run('Dintr-o privire, mesajul.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))
    # garda em-dash
    full = '\n'.join(pp.text for pp in doc.paragraphs)
    for ch in ['—', '–']:
        if ch in full:
            print('  ATENTIE em/en-dash in FILM!')
    import re as _re
    print('  leftover C14/Compunere:', len(_re.findall(r'C14|Compunere|COMPUNERE', full)))


if __name__ == '__main__':
    main()
