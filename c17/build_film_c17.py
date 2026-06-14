#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c17.py - FILM-Excel-17-Sistematizare.docx (script video procedural).

Structura = c16 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C17 SISTEMATIZARE.
Garda C17: a face munca reluabila. ZERO automatizare ca identitate (=C18),
ZERO reguli/praguri (=C19), ZERO ownership (=C20). Garda OGLINDA. Anti-SOP.
Cifre in Excel (R-V02.15). Zero em-dash.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'c17/FILM-Excel-17-Sistematizare.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C17 SISTEMATIZARE')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 5 AUTONOMIE · Construcția 17 din 20 · deschide treapta T5')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK->DEMO->CONTROL->REVEAL.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('Munca ta merge perfect. Până pleci tu. Astăzi nu mai facem raportul, facem munca reluabilă: un sistem pe care altcineva îl pornește fără tine.')
    lbl('MIZA')
    txt('Câtă vreme munca depinde de o singură persoană, e un risc ascuns: în concediu, la '
        'demisie, la suprasolicitare, se oprește. Trecem de la o muncă pe care o știai doar tu '
        'la o hartă vie în workbook, pe care un om instruit o pornește fără tine. Sistematizarea '
        'face munca să continue cu oricine, nu doar cu autorul ei.')
    lbl('AHA (INSIGHT)')
    txt('O muncă făcută de două ori nu mai e o livrare. E un sistem ascuns în workbook.')
    lbl('MANTRA · TRAINITY')
    txt('Nu o fac din nou. O fac sistem.')
    lbl('WOW (PAYOFF FINAL)')
    txt('Un proces care trăia doar în capul tău. Acum o hartă vie în workbook, pe care altcineva o pornește fără tine.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Pleci, și munca o reia altcineva.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion. FRAZĂ DE IMPACT = exec-phrase.')
    execs = [
        ('1 · REALITATE', 'Munca pe care o știi doar tu',
         'Munca merge, dar numai cu tine la cârmă. Fără tine, nu pornește: tu ești ocazia.'),
        ('2 · INVENTAR', 'Componentele reale',
         'Foile reale, prin link și oglindă. Harta nu descrie munca, o arată vie.'),
        ('3 · SURSĂ', 'Un singur adevăr per intrare',
         'Sursa canonică, fixată ca SRC_, cu oglinda ei. Nimeni nu mai întreabă care fișier.'),
        ('4 · RELUARE', 'Ordinea legată de obiecte',
         'Pașii, un lanț navigabil spre obiecte reale. Repetitivii, marcați candidat C18, nu automatizați.'),
        ('5 · EXTERNALIZARE', 'Cunoașterea scoasă din cap',
         'Convențiile din cap devin parametri vii. Cine reia e un profil, nu un nume.'),
        ('6 · SUBSTITUT', 'Reluabil fără autor',
         'Altcineva pornește din _SISTEM și reia munca fără tine. Apoi predarea spre automatizare.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C17')
    txt('Construcția 17 SISTEMATIZARE deschide treapta de autonomie. Treptele precedente au dat '
        'un răspuns, l-au făcut raport și l-au predat gata de decizie. Dar raportul l-ai făcut tu, '
        'o dată. În concediu, munca se oprește: tu ești ocazia, fără tine nu pornește. Acum nu mai '
        'facem raportul, facem munca reluabilă.')
    txt('Pentru prima dată, obiectul nu mai sunt datele, nici raportul, ci munca însăși. O scriem '
        'ca hartă vie în workbook, legată de foile reale, pe care un om instruit o pornește fără '
        'noi. C17 face munca reluabilă de altcineva; să ruleze singură începe mai sus.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Inventariază componentele reale ale workbook-ului și propune ordinea pașilor de reluare, '
        'legați de obiecte. Marchează ce ar fi candidat de automatizare, dar nu automatizează '
        'nimic și nu scrie reguli; aceea rămâne la treapta următoare.')
    lbl('EXCEL · _SISTEM')
    txt('Găzduiește foaia _SISTEM: componentele prin HYPERLINK și COUNTA, sursele ca SRC_, pașii '
        'legați de OBJ_, convențiile ca PARAM_, START AICI și punctul de reluare. Totul prin '
        'legături vii care se rup scoase din workbook. Valorile sursă rămân neatinse.')
    lbl('CONTROL UMAN')
    txt('Scriem harta, alegem sursa canonică, fixăm ordinea, scoatem convențiile din cap în '
        'parametri, facem testul substitutului. Nu automatizăm pașii și nu punem reguli de '
        'control: facem munca reluabilă, nu o punem să ruleze singură.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CELE 6 FENOMENE')
    fenomene = [
        ('01 · CUNOAȘTERE TRIBALĂ',
         'Procesul trăiește doar în capul tău, nescris. Corecția: îl scrii ca hartă vie în _SISTEM, legată de componentele reale.'),
        ('02 · RELUARE DIN MEMORIE',
         'De fiecare dată reconstruiești pașii din minte, cu variații. Corecția: ordinea de reluare se scrie o dată, ca lanț legat de obiecte.'),
        ('03 · SURSĂ AMBIGUĂ',
         'Aceeași intrare vine din două locuri. Corecția: o singură sursă de adevăr per intrare, fixată ca named range SRC_.'),
        ('04 · CONVENȚII NESCRISE',
         'Alegerile tale nu sunt nicăieri, doar tu le știi. Corecția: convențiile devin parametri vii, named ranges PARAM_.'),
        ('05 · DOAR TU POȚI PORNI',
         'Nimeni nu reia fără să te întrebe de unde începe. Corecția: scrii profilul de competență care poate relua, nu numele tău.'),
        ('06 · GÂTUIREA PE AUTOR',
         'Pleci, munca se oprește. Corecția: harta reluabilă scoate munca din tine, fără să o automatizeze încă.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 6 fenomene care țin munca legată de autor, fiecare cu corecția lui reluabilă. '
        'Munca e pregătită să devină un sistem.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · REALITATE',
            'obiectiv': 'Constatăm că munca merge dar numai cu autorul, încercăm s-o reia altcineva și se blochează, creăm foaia _SISTEM și declarăm harta reluabilă.',
            'stare': 'Atenție. Un workbook care merge perfect, dar numai cu tine la cârmă.',
            'voce': '„Etapa 1: realitate. Munca iese de fiecare dată. Dar ca s-o reiei, trebuie să știi din cap de unde pleci. E muncă, nu sistem."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C17, foile împrăștiate; rugăm un coleg să reia și se blochează la prima întrebare.',
            'ai': 'AI-ul confirmă că procesul există în workbook, dar harta lipsește. Noi decidem să o scriem.',
            'control': 'Punem regula-test: scriem o hartă vie, nu o procedură de citit. „O muncă făcută de două ori e un sistem care așteaptă să fie scris."',
            'reveal': 'Pe ecran: foaia _SISTEM, goală, fixată prima, gata de construit.',
        },
        {
            'nume': 'ETAPA 2 · INVENTAR',
            'obiectiv': 'Inventariem componentele reale ale workbook-ului si le facem navigabile in blocul A, prin HYPERLINK si COUNTA. Promptul 1.',
            'stare': 'Curiozitate metodică. Vedem din ce e făcută munca, nu ce raport iese.',
            'voce': '„Etapa 2: inventar. Ce foi sunt surse, ce e lucru, ce e ieșire. Componentele există deja; le facem vizibile și navigabile."',
            'demo': 'Pe ecran: rulăm Promptul 1; AI-ul propune inventarul componentelor; construim blocul A cu link viu și COUNTA.',
            'ai': 'AI-ul inventariază. Noi confirmăm rolurile; nu lăsăm AI-ul să automatizeze sau să decidă final.',
            'control': 'Arătăm garda OGLINDĂ pe viu: „COUNTA arată starea, nu o schimbă. Harta nu rulează munca, o reflectă."',
            'reveal': 'Pe ecran: blocul A, componentele reale prin link, fiecare cu cât conține, live.',
        },
        {
            'nume': 'ETAPA 3 · SURSĂ',
            'obiectiv': 'Fixăm o singură sursă de adevăr per intrare, ca named range SRC_, eliminăm dublura, punem oglinda ROWS în blocul B.',
            'stare': 'Precizie. O hartă reluabilă cere un singur adevăr per intrare.',
            'voce': '„Etapa 3: sursă. Pentru fiecare intrare, o singură sursă canonică, fixată ca SRC_. Reluarea nu mai întreabă care fișier."',
            'demo': 'Pe ecran: definim SRC_ pe zona reală, eliminăm dublura, punem ROWS(SRC_) ca oglindă.',
            'ai': 'AI-ul semnalează intrările cu surse multiple. Noi declarăm una canonică; ambiguitatea dispare.',
            'control': 'Confirmăm garda OGLINDĂ: „ROWS doar arată câte rânduri are sursa. Nu verifică un prag și nu blochează nimic."',
            'reveal': 'Pe ecran: blocul B, fiecare intrare cu un singur SRC_ și oglinda lui vie.',
        },
        {
            'nume': 'ETAPA 4 · RELUARE',
            'obiectiv': 'Scriem ordinea de reluare ca lanț navigabil legat de obiecte, în blocul C, cu rezultatul oglindit; marcăm pașii repetitivi candidat C18, fără să automatizăm. Promptul 2.',
            'stare': 'Construcție calmă. Pașii devin o hartă navigabilă, nu un motor.',
            'voce': '„Etapa 4: reluare. Fiecare pas duce la obiectul real și arată ce iese. Harta spune ce se face și unde, nu o face în locul tău."',
            'demo': 'Pe ecran: rulăm Promptul 2; AI-ul propune lanțul de pași; legăm fiecare prin HYPERLINK la OBJ_, oglindim rezultatul.',
            'ai': 'AI-ul așază ordinea și marchează candidații de automatizare. Noi îi lăsăm manuali; nu scriem niciun macro.',
            'control': 'Ținem granița spre C18: „Pasul repetitiv e candidat C18, dar aici rămâne acțiune de om. Nu automatizăm încă."',
            'reveal': 'Pe ecran: blocul C, lanțul de pași navigabil, cu candidat C18 marcat și neautomatizat.',
        },
        {
            'nume': 'ETAPA 5 · EXTERNALIZARE',
            'obiectiv': 'Scoatem convențiile din cap ca named ranges PARAM_ și scriem profilul de competență care poate relua, în blocurile D și E.',
            'stare': 'Eliberare. Ce știai doar tu devine vizibil în workbook.',
            'voce': '„Etapa 5: externalizare. Alegerile pe care le făceai că-așa-știu-eu devin parametri vii. Cine poate relua e o competență, nu un nume."',
            'demo': 'Pe ecran: scriem PARAM_ cu valorile lor, legăm pașii de ele; scriem profilul de competență în blocul D.',
            'ai': 'AI-ul scoate la iveală alegerile implicite din pași. Noi le scriem ca PARAM_, vizibile.',
            'control': 'Ținem granița spre C20: „Spunem cine poate relua, nu cine deține. Nu predăm răspunderea și nu numim un titular."',
            'reveal': 'Pe ecran: blocurile D și E, profilul și parametrii, cunoașterea tribală ieșită din cap.',
        },
        {
            'nume': 'ETAPA 6 · SUBSTITUT',
            'obiectiv': 'Punem START AICI și punctul de reluare, facem testul substitutului pe viu, marcăm granița C17 spre C18 și predăm harta.',
            'stare': 'Claritate. Proba finală: altcineva pornește din foaie, fără tine.',
            'voce': '„Etapa 6: substitut. Dau workbook-ul unui coleg și plec din cameră. Pornește din START AICI și reia munca doar din foaie."',
            'demo': 'Pe ecran: colegul pornește din _SISTEM, urmează linkurile, atinge obiectele, reia munca; testăm și anti-SOP, scoțând o foaie.',
            'ai': 'AI-ul nu mai intervine: harta se susține singură. Dacă substitutul se blochează pe ceva nescris, mai avem de scris.',
            'control': 'Confirmăm anti-SOP și handoff: „Copiezi _SISTEM în Word, se rupe. C17 face munca reluabilă de un om; C18 scoate omul din pașii repetitivi."',
            'reveal': 'Pe ecran: substitutul reia munca fără autor; harta predată spre treapta de automatizare.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · MUNCA DEVENITĂ SISTEM')
    txt('C17 face munca reluabilă de altcineva. Harta _SISTEM e vie, legată de foile reale, se '
        'rupe scoasă din workbook: nu e o procedură, e un sistem. Un om instruit o pornește din '
        'START AICI și reia munca fără autor. C17 face munca reluabilă; automatizarea pașilor '
        'repetitivi, guvernarea prin reguli și predarea proprietății încep la treptele următoare. '
        'Predăm o hartă cu pașii candidat C18 marcați, cu valorile sursă neatinse și suma '
        'conservată cap-coadă.')
    p = doc.add_paragraph()
    r = p.add_run('Pleci, și munca o reia altcineva.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))


if __name__ == '__main__':
    main()
