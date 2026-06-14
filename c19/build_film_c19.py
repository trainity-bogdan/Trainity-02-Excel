#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c19.py - FILM-Excel-19-Guvernare.docx (script video procedural).

Structura = c17/c18 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C19 GUVERNARE.
Axa CONTROL / GUVERNARE: reguli, praguri LIMIT_, validare la sursa, semnale,
STATUS OK/ATENTIE/OPRIT, exceptii, fail-safe, oprire controlata.
Garda C19: a tine sistemul corect prin reguli. ZERO automatizare ca identitate (=C18),
ZERO ownership/responsabil/escaladare (=C20), ZERO dashboard de citit (=T4), ZERO
monitorizare/babysitting. Verb GUVERNEZ. Artefact _GUVERNARE.
Cifre in Excel (R-V02.15). Zero em-dash, zero en-dash.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'c19/FILM-Excel-19-Guvernare.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C19 GUVERNAREA SISTEMULUI PRIN REGULI')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 5 AUTONOMIE · Construcția 19 din 20')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK->DEMO->CONTROL->REVEAL în interior.')
    txt('Arc: input greșit -> regulă -> prag și stare -> excepție -> oprire controlată -> testul ochilor închiși.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('Cum se ține corect fără ochiul meu? Motorul tău rulează, lanțul se mișcă fără mâna ta. '
        'Și totuși, pe o intrare greșită, produce un rezultat greșit în tăcere, iar tu tot '
        'verifici „a ieșit bine?" la fiecare ciclu. Astăzi nu mai supraveghem munca. O guvernăm '
        'prin reguli.')
    lbl('BOMBĂ')
    txt('Motorul tău rulează. Dar tot tu verifici că n-a greșit, de fiecare dată.')
    lbl('MIZA')
    txt('Câtă vreme cineva trebuie să confirme de fiecare dată că rezultatul e corect, sistemul '
        'nu e autonom: atenția unui om e încă plafonul, iar o intrare greșită nedetectată intră '
        'în decizii ca și cum ar fi bună. Guvernarea prin reguli mută verificarea din mintea '
        'omului în sistem: deviațiile previzibile sunt prinse, semnalate sau blocate înainte să '
        'producă pagubă tăcută.')
    lbl('GREȘEALA CLASICĂ')
    txt('Confunzi «merge» cu «merge corect». Un sistem care rulează nu e același lucru cu un '
        'sistem în care poți avea încredere.')
    lbl('AHA (INSIGHT)')
    txt('Un sistem în care ai încredere nu e cel pe care îl urmărești. E cel care se prinde '
        'singur când greșește.')
    lbl('MANTRA · TRAINITY')
    txt('Nu o supraveghez. O guvernez prin reguli.')
    lbl('WOW (PAYOFF FINAL)')
    txt('Un sistem care mergea doar cât stăteai cu ochii pe el. Acum, pe o intrare greșită, se '
        'oprește singur și aprinde semnalul, fără tine.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Pleci, și munca se ține singură.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion. FRAZĂ DE IMPACT = exec-phrase.')
    execs = [
        ('1 · REALITATE', 'Plafonul ascuns',
         'Motorul merge. Astăzi am descoperit că tot eu verific dacă a mers corect.'),
        ('2 · DEVIAȚII', 'Privirea înainte de regulă',
         'Nu întreb ce a ieșit. Întreb ce poate devia previzibil și ar curge tăcut.'),
        ('3 · REGULA', 'Intrarea respinsă la sursă',
         'Nu am citit eu rezultatul. Regula respinge intrarea greșită înainte să intre în motor.'),
        ('4 · PRAG-SEMNAL', 'Semnalul care acționează',
         'O valoare derapează, STATUS se schimbă singur. Semnalul nu se citește, acționează.'),
        ('5 · OPRIRE', 'Oprirea controlată',
         'Pe o eroare gravă, fail-safe-ul reține rezultatul corupt. Se oprește singur, fără mâna mea.'),
        ('6 · TEST', 'Ochii închiși',
         'Am plantat deviații și am plecat. Regulile au prins, au semnalat, au oprit fără mine.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C19')
    txt('Construcția 19 GUVERNAREA continuă treapta de autonomie. Construcția dinainte a scos '
        'execuția din mâinile autorului: motorul rulează, lanțul se mișcă fără el. Funcționează. '
        'Și totuși, pe o intrare greșită, produce un rezultat greșit în tăcere, iar autorul tot '
        'verifică „a ieșit bine?" la fiecare ciclu. Atenția lui a devenit plafonul.')
    txt('Problema nu este că motorul lipsește. Problema este că „merge" nu înseamnă „merge corect". '
        'Un motor pe care tot tu îl verifici te-a scos din execuție, nu din grijă. Acum nu mai '
        'supraveghem munca, o guvernăm prin reguli. C19 nu repară motorul și nu îl pornește: ia un '
        'sistem care rulează și îi dă regulile care îl țin corect. C19 face controlul, nu mișcarea: '
        'motorul vine de jos, predarea proprietății începe mai sus.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Listează, pentru fiecare coloană de intrare, valorile întâlnite și intervalul permis, și '
        'propune regula de validare la sursă și logica semnalului: ce abatere aprinde ATENȚIE și ce '
        'eroare trece STATUS în OPRIT. Propune pragurile. Nu corectează automat nicio valoare, nu '
        'configurează execuția lanțului și nu desemnează cine răspunde. AI propune, omul decide.')
    lbl('EXCEL · _GUVERNARE')
    txt('Găzduiește foaia _GUVERNARE: plicul ca praguri vii LIMIT_, validarea la sursă prin Data '
        'Validation, semnalele care calculează STATUS din reguli, lista de excepții și fail-safe-ul '
        'OUTPUT_GUVERNAT legat de STATUS. Totul prin formule vii care recalculează singure la fiecare '
        'schimbare în date. Valorile sursă rămân neatinse, suma e conservată cap-coadă.')
    lbl('CONTROL UMAN')
    txt('Scriem registrul deviațiilor, fixăm pragurile, punem validarea la sursă, decidem unde e '
        'linia dintre ATENȚIE și OPRIT, legăm fail-safe-ul de STATUS, facem testul ochilor închiși. '
        'Marcăm excepția, dar nu desemnăm responsabilul și nu escaladăm la nimeni: guvernăm '
        'sistemul, nu predăm proprietatea. Aceea începe la treapta următoare.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CE PARE CONTROL DAR ESTE DOAR OCHIUL TĂU')
    fenomene = [
        ('01 · BABYSITTING',
         'Stai cu ochii pe sistem ca să-l prinzi când greșește; monitorizarea te ține prezent. '
         'Corecția: regulile prind în locul ochiului, sistemul se ține corect fără tine.'),
        ('02 · DASHBOARD DE CITIT',
         'Un panou frumos care afișează stări pentru un om care decide; acela e T4. '
         'Corecția: semnalul C19 acționează, schimbă starea și leagă oprirea, nu se citește.'),
        ('03 · DOAR REFRESH',
         'Aduci setul la zi și crezi că l-ai ținut corect; acela e refresh, nu guvernare. '
         'Corecția: regulile țin corect un sistem care rulează, nu reîmprospătează un set.'),
        ('04 · MOTORUL CONFUNDAT CU CONTROLUL',
         'Lanțul rulează singur, dar pe o intrare proastă produce gunoi tăcut; mișcarea vine de jos. '
         'Corecția: C19 nu pune lanțul în mișcare, îi adaugă regulile care îl țin corect.'),
        ('05 · RESPONSABILUL DEGHIZAT',
         'Întrebi cine deține și cine răspunde la o abatere; aceea e predarea proprietății, mai sus. '
         'Corecția: _GUVERNARE marchează și oprește, dar nu scrie un nume și nu escaladează.'),
        ('06 · VERIFIC LA FINAL',
         'Citești rezultatul la capăt și speri că e bun; greșeala se vede târziu, după pagubă. '
         'Corecția: regula prinde abaterea în momentul apariției, la sursă, nu la capăt.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 6 fenomene care pretind control, dar țin sistemul legat de ochiul autorului, '
        'fiecare cu corecția lui prin reguli. Recunoaște-le, ca să guvernezi prin reguli, nu să '
        'patrulezi sistemul.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · Sistemul care merge doar cât te uiți',
            'sub': 'REALITATE',
            'obiectiv': 'Pornim motorul moștenit ca de obicei, vedem cum o intrare greșită produce gunoi în '
                        'tăcere, creăm foaia _GUVERNARE și declarăm că o guvernăm prin reguli.',
            'stare': 'Recunoaștere acută. Motorul merge fără mâna mea. Și tot eu verific dacă a mers corect.',
            'voce': '„Motorul rulează, lanțul se mișcă singur. Și totuși, dacă în date intră o valoare '
                    'imposibilă, o însumă liniștit și scoate un rezultat greșit, fără să clipească. Merge. '
                    'Nu înseamnă că merge corect."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C19, motorul moștenit rulează singur; strecurăm o '
                    'valoare imposibilă și o vedem însumată tăcut într-un rezultat greșit.',
            'ai': 'Încă nu intervine AI. Doar privim cum o intrare greșită curge prin motor fără să fie prinsă.',
            'control': 'Punem regula-test: „Verificarea iese din mintea mea și intră în sistem. Nu o mai '
                       'supraveghez, o guvernez prin reguli."',
            'reveal': 'Pe ecran: foaia _GUVERNARE, fixată ca foaia controlului. Aici nu reparăm motorul, '
                      'îi dăm regulile care îl țin corect.',
        },
        {
            'nume': 'ETAPA 2 · Ce poate să devieze previzibil',
            'sub': 'DEVIAȚII',
            'obiectiv': 'Trecem în revistă coloanele de intrare, cerem Copilot să listeze valorile și '
                        'intervalul permis, scriem plicul așteptat ca praguri vii LIMIT_. Promptul 1.',
            'stare': 'Privire metodică înainte de regulă. Nu întreb ce a ieșit, întreb ce poate devia previzibil.',
            'voce': '„Înainte de orice regulă, văd ce poate strica rezultatul. Ce coloană poate primi o '
                    'valoare imposibilă și care e plicul în care o intrare e încă plauzibilă."',
            'demo': 'Pe ecran: deschidem Copilot pe _GUVERNARE, lipim Promptul 1; AI propune, pe fiecare '
                    'coloană, valorile și intervalul permis; scriem pragurile ca named ranges LIMIT_.',
            'ai': 'AI parcurge fiecare coloană de intrare și propune intervalul permis plus o regulă de '
                  'validare la sursă. Nu corectează nicio valoare și nu desemnează pe nimeni. AI propune, omul decide.',
            'control': 'Ținem garda OGLINDĂ pe praguri: „Pragul marchează, nu repară. Spune «în afara '
                       'plicului», nu schimbă valoarea în locul meu. E un parametru viu, citit de fiecare regulă."',
            'reveal': 'Pe ecran: plicul așteptat scris ca praguri vii LIMIT_, în celulă. Schimbi pragul, '
                      'regula se mută cu el.',
        },
        {
            'nume': 'ETAPA 3 · Regula care prinde intrarea greșită',
            'sub': 'REGULA',
            'obiectiv': 'Punem Data Validation la sursă legată de praguri, cerem Copilot logica semnalului, '
                        'construim STATUS OK / ATENȚIE / OPRIT calculat din reguli. Promptul 2.',
            'stare': 'Precizie. Plicul devine reguli care acționează: intrarea greșită e respinsă la sursă.',
            'voce': '„Pe zona de intrare pun Data Validation legată de praguri. O valoare imposibilă e '
                    'respinsă la tastare, nu ajunge niciodată în motor. Regula prinde la sursă, nu la capăt."',
            'demo': 'Pe ecran: legăm Data Validation de LIMIT_; lipim Promptul 2; AI propune ce abatere '
                    'aprinde ATENȚIE și ce eroare trece STATUS în OPRIT; construim STATUS_SISTEM ca formulă peste semnale.',
            'ai': 'AI așază logica semnalului peste praguri: ce e doar suspect și ce e grav. Noi corectăm '
                  'pragurile și decidem unde e linia. AI nu configurează execuția lanțului și nu se rezumă la un refresh.',
            'control': 'Confirmăm că STATUS nu e o etichetă: „STATUS nu e o eticheta pe care o pun eu, e o '
                       'stare care se recalculează singură la fiecare schimbare în date."',
            'reveal': 'Pe ecran: intrarea greșită respinsă la sursă, pragul citit de reguli, STATUS care se '
                      'recalculează singur, colorat verde, galben, roșu.',
        },
        {
            'nume': 'ETAPA 4 · Pragul și semnalul',
            'sub': 'PRAG-SEMNAL',
            'obiectiv': 'Bagăm o intrare greșită și o vedem respinsă la sursă; o valoare derapează intern și '
                        'pragul aprinde ATENȚIE, STATUS se schimbă singur; facem testul anti-dashboard.',
            'stare': 'Calm controlat. Semnalul nu se mulțumește să afișeze. Acționează, schimbă starea.',
            'voce': '„Întrebarea care separă guvernarea de un panou: semnalul doar arată o culoare, sau face '
                    'ceva? Aici semnalul schimbă STATUS și, prin el, leagă oprirea lanțului."',
            'demo': 'Pe ecran: tastăm o cantitate în afara intervalului, Data Validation o oprește pe loc; '
                    'apoi o agregare iese din plic, STATUS trece din OK în ATENȚIE fără să atingem nimic.',
            'ai': 'AI nu mai intervine. Regulile observă singure: nu omul vede abaterea și schimbă eticheta, '
                  'regula o observă și starea se mută singură.',
            'control': 'Ținem garda anti-T4: „Dacă ar fi doar un dashboard de citit pentru un om care decide, '
                       'n-ar fi guvernare. Semnalul acționează: schimbă STATUS și leagă mai departe oprirea."',
            'reveal': 'Pe ecran: intrarea greșită prinsă la sursă, STATUS schimbat singur de la o derapare '
                      'internă. Semnalul acționează, nu se citește.',
        },
        {
            'nume': 'ETAPA 5 · Excepția și oprirea controlată',
            'sub': 'OPRIRE',
            'obiectiv': 'Cazul neprevăzut cade în lista de excepții fără să treacă tăcut; construim '
                        'OUTPUT_GUVERNAT legat de STATUS=OPRIT; marcăm granița spre predarea proprietății.',
            'stare': 'Onestitate operațională. Excepția nu se ascunde. Pe OPRIT, sistemul se oprește singur.',
            'voce': '„O categorie necunoscută, un caz pe care nicio regulă nu îl prevedea, în loc să treacă '
                    'tăcut cade în lista de excepții, marcat «de privit». Excepția nu se ascunde și nu se autocorectează."',
            'demo': 'Pe ecran: o categorie necunoscută cade în lista de excepții din _GUVERNARE; construim '
                    'OUTPUT_GUVERNAT: pe STATUS=OPRIT rezultatul e reținut, altfel trece totalul validat.',
            'ai': 'AI nu intervine. Decizia de unde e linia e judecată umană; oprirea însăși e automată, '
                  'legată de stare, nu o apasă nimeni.',
            'control': 'Ținem granița spre treapta următoare: „_GUVERNARE prinde abaterea, aprinde semnalul '
                       'și oprește lanțul, dar nu scrie un nume lângă excepție și nu escaladează. Spune «aici e o '
                       'abatere și sistemul s-a oprit», nu «tu răspunzi de ea». Cine deține și cine preia e treaba '
                       'treptei următoare."',
            'reveal': 'Pe ecran: lista de excepții cu cazul scos la lumină și OUTPUT_GUVERNAT pe OPRIT, '
                      'reținând rezultatul corupt. Oprirea e automată, fail-safe, nu o apeși tu.',
        },
        {
            'nume': 'ETAPA 6 · Testul ochilor închiși',
            'sub': 'TEST',
            'obiectiv': 'Plantăm deviații și plecăm din fața ecranului, regulile prind, semnalează și opresc '
                        'singure; marcăm că STATUS=OPRIT e intenționat; predăm sistemul guvernat.',
            'stare': 'Predare conștientă. Proba finală: deviațiile plantate sunt prinse fără ochiul meu pe ecran.',
            'voce': '„Strecor în date câteva valori imposibile și un caz neprevăzut, apoi plec din fața '
                    'ecranului. Fără mine, regulile prind anomaliile, semnalul trece STATUS în OPRIT, fail-safe-ul '
                    'reține rezultatul."',
            'demo': 'Pe ecran: în demonstrația tehnică, datele vin cu deviații deja plantate. STATUS=OPRIT '
                    'este intenționat, dovada că sistemul s-a prins singur, nu un defect. _GUVERNARE prinde '
                    'anomaliile pe care motorul moștenit le-ar fi lăsat să curgă tăcut.',
            'ai': 'AI nu mai intervine: regulile se susțin singure. Dacă o deviație trece tăcut, mai e o '
                  'regulă de scris; sistemul nu se ține încă singur.',
            'control': 'Rostim semnătura și granița: „Nu mai patrulez sistemul ca să-l prind când greșește. '
                       'I-am scris regulile, pragurile, semnalele și oprirea. Sistemul se ține corect singur, încă '
                       'deținut de mine. Cine îi preia proprietatea începe la C20."',
            'reveal': 'Pe ecran: sistemul prinde, semnalează și oprește fără autor; predat ca sistem guvernat, '
                      'reguli, praguri, semnale și oprire vii în _GUVERNARE, spre predarea proprietății la C20.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('SUBTITLU'); txt(e['sub'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · SISTEMUL CARE SE ȚINE SINGUR')
    txt('C19 face sistemul să se țină corect singur, fără ochiul autorului. Foaia _GUVERNARE e '
        'vie: praguri LIMIT_, validare la sursă, semnale care calculează STATUS, listă de excepții '
        'și un fail-safe care leagă rezultatul de stare. Intrarea greșită e respinsă la sursă, '
        'excepția nu trece tăcut, iar pe o eroare gravă STATUS devine OPRIT și OUTPUT_GUVERNAT '
        'reține rezultatul corupt. STATUS=OPRIT din demonstrație este intenționat: dovada că '
        'sistemul se prinde singur. C19 nu construiește motorul (vine de jos) și nu desemnează '
        'responsabilul (granița de sus): autorul rămâne titularul regulilor. Predăm un sistem '
        'guvernat, cu valorile sursă neatinse și suma conservată cap-coadă, spre predarea '
        'proprietății la treapta următoare.')
    p = doc.add_paragraph()
    r = p.add_run('Pleci, și munca se ține singură.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))


if __name__ == '__main__':
    main()
