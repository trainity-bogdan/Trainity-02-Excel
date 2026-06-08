#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c14.py - FILM-Excel-14-Compunere.docx (script video procedural).

Structura = c13 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C14 COMPUNERE.
Garda C14: a asaza spatial. Pagina cu focar/ierarhie, ZERO obiect singular (C13),
zero mesaj (C15), zero livrare (C16). Cifre in Excel (R-V02.15).
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = 'c14/FILM-Excel-14-Compunere.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    os.makedirs('c14', exist_ok=True)
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C14 COMPUNERE')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 4 RAPORTARE · Construcția 14 din 20 · a doua din T4')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK->DEMO->CONTROL->REVEAL.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('Graficele sunt corecte. Pagina te încurcă. Astăzi așezăm obiectele ca ochiul să prindă întâi ce contează.')
    lbl('MIZA')
    txt('Eșecul silențios al paginii prost compuse: o decizie lentă sau greșită luată pe o pagină '
        'care pare completă, iar vina cade pe cititor, nu pe autor. Trecem de la obiecte oneste, '
        'împrăștiate fără ordine, la o pagină de raport cu focar, ierarhie și traseu de citire, '
        'în care decidentul prinde decizia dintr-o privire.')
    lbl('AHA (INSIGHT)')
    txt('Aceleași grafice, altă ordine, altă decizie.')
    lbl('MANTRA · TRAINITY')
    txt('Compun privirea, nu pagina.')
    lbl('WOW (PAYOFF FINAL)')
    txt('Aceleași grafice, aceleași cifre, altă așezare. Decizia se schimbă mutând blocurile, nu conținutul.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Ce vede ochiul întâi schimbă decizia.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion. FRAZĂ DE IMPACT = exec-phrase.')
    execs = [
        ('1 · ZIDUL', 'Pagina fără ordine',
         'Ai obiecte corecte și o pagină fără un întâi. Ochiul aterizează la întâmplare și rătăcește.'),
        ('2 · POZIȚIA', 'Poziția e importanță',
         'Unde așezi un obiect spune cât contează. Poziția nu e neutră, e un argument despre importanță.'),
        ('3 · TRASEUL', 'Ordinea de citire',
         'Un focar atins primul, apoi un traseu: ce vede ochiul după și ce coboară din prim-plan.'),
        ('4 · IERARHIA', 'Greutate inegală',
         'Mărime, poziție, contrast, spațiu: importanța se vede instant, nu se caută.'),
        ('5 · SPAȚIUL', 'Compoziția guvernată',
         'Spațiul gol izolează focarul. Toată pagina se aliniază la răspunsul venit din analiză.'),
        ('6 · PROBA', 'Al doilea ochi',
         'Un cititor de trei secunde prinde aceeași prioritate. Pagina e gata de predat pentru mesaj.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C14')
    txt('Construcția 14 COMPUNERE preia obiectele vizuale oneste de la treapta de vizualizare. '
        'Avem grafice corecte, fiecare adevărat singur. Dar pe o pagină fără ordine, adevărul '
        'există peste tot și nu se vede nicăieri. Acum nu mai desenăm niciun obiect nou; așezăm '
        'obiectele pe pagină.')
    txt('Pentru prima dată, ordinea nu mai e un detaliu. E o decizie: aceleași grafice, altă '
        'așezare, altă decizie. C14 alege ce vede ochiul întâi.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Propune o ierarhie și o ordine de citire pentru obiectele date și o variantă de '
        'compoziție. Aranjează repede. Nu garantează că ordinea servește decizia; aceea o pune omul.')
    lbl('EXCEL · COMPUNERE')
    txt('Găzduiește foaia Compunere: obiectele primite, focarul, ierarhia, traseul de citire, '
        'gruparea, spațiul și un before/after de rearanjare. Valorile sursă rămân neatinse.')
    lbl('CONTROL UMAN')
    txt('Stabilim focarul după răspunsul venit din analiză, ordonăm traseul, retrogradăm '
        'secundarul, facem ierarhia și grupăm. Nu desenăm obiectul și nu formulăm mesajul: așezăm pagina.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CELE 6 FENOMENE')
    fenomene = [
        ('01 · PAGINA-ZID',
         'Zeci de obiecte de aceeași mărime, niciunul mai important, ochiul nu știe unde să meargă. Operația: stabilești un singur focar, atins primul.'),
        ('02 · ESENȚIALUL ÎNGROPAT',
         'Cel mai important rezultat stă jos, în colț, la fel de mare ca un detaliu. Operația: esențialul urcă în primul punct de contact al ochiului.'),
        ('03 · ORDINEA DE PRODUCȚIE',
         'Pui obiectele în ordinea în care le-ai făcut, nu în care trebuie văzute. Operația: ordinea de citire guvernată de răspuns, nu de cronologie.'),
        ('04 · GREUTATE EGALĂ',
         'Titlu, total, detaliu, notă, toate la fel de mari: nicio ierarhie. Operația: greutate inegală prin mărime, poziție, contrast, spațiu.'),
        ('05 · PROXIMITATEA CARE MINTE',
         'Lucruri legate stau departe, lucruri nelegate stau lipite. Operația: grupezi spațial ce ține împreună, separi ce nu e legat.'),
        ('06 · HORROR VACUI',
         'Umpli orice spațiu gol cu încă un obiect. Operația: spațiul gol ca instrument de ierarhie, izolează focarul și separă grupurile.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 6 fenomene de așezare, fiecare cu operația lui de compunere. Obiectele sunt '
        'pregătite să intre într-o pagină coerentă.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · ZIDUL',
            'obiectiv': 'Preluăm obiectele vizuale oneste de la treapta de vizualizare, constatăm că o pagină fără ordine nu are un întâi și vedem unde cade ochiul primul.',
            'stare': 'Atenție. Obiecte corecte, o pagină-zid: ochiul aterizează la întâmplare și rătăcește.',
            'voce': '„Etapa 1: zidul. Avem grafice corecte. Dar pe pagină, toate la fel, niciun focar. Ochiul nu știe unde să se uite. Îi dăm o ordine."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C14, foaia Compunere; arătăm obiectele așezate în ordinea producției, fără focar.',
            'ai': 'AI-ul arată că obiectele există deja, corecte. Noi observăm că pagina nu are un întâi, fără să redesenăm nimic.',
            'control': 'Punem regula-test: așezăm, nu redesenăm. „Obiecte corecte pe o pagină fără ordine nu fac încă o decizie. Le dăm un traseu."',
            'reveal': 'Pe ecran: pagina-zid, cu obiectele corecte dar fără un punct de start.',
        },
        {
            'nume': 'ETAPA 2 · POZIȚIA',
            'obiectiv': 'Descoperim că poziția semnalează importanță și alegem ce întâlnește ochiul primul. Promptul 1.',
            'stare': 'Atenție la poziție. Înainte să așezăm, vedem că poziția e un argument, nu un detaliu.',
            'voce': '„Etapa 2: poziția. Ce stă sus, mare sau în centru spune important. Poziția nu e neutră, e un argument. Importanța o așezăm noi."',
            'demo': 'Pe ecran: mutăm același obiect sus, jos, în colț; rulăm Promptul 1; AI-ul propune ierarhia și ordinea de citire.',
            'ai': 'AI-ul propune ce vede ochiul întâi, ce după, ce se retrogradează. Noi judecăm dacă ordinea servește decizia.',
            'control': 'Confirmăm că poziția comunică importanță. „Unde așezăm un obiect spune cât contează, vrând-nevrând. De acum o folosim intenționat."',
            'reveal': 'Pe ecran: același obiect, trei poziții, trei niveluri de importanță citită.',
        },
        {
            'nume': 'ETAPA 3 · TRASEUL',
            'obiectiv': 'Stabilim focarul atins primul, ordonăm traseul de citire și retrogradăm ce nu mută decizia.',
            'stare': 'Construcție calmă. Pagina capătă un început și un drum.',
            'voce': '„Etapa 3: traseul. Un singur focar, obiectul care poartă răspunsul, atins primul. Apoi al doilea, al treilea. Ce nu mută decizia coboară din prim-plan."',
            'demo': 'Pe ecran: plasăm focarul sus-stânga, construim fluxul focar-suport-detaliu, retrogradăm un tabel secundar.',
            'ai': 'AI-ul sugerează traseul. Noi confirmăm că focarul poartă răspunsul și că secundarul a coborât, nu a dispărut.',
            'control': 'Aplicăm filtrul: fără acest element în prim-plan, decidentul hotărăște la fel? Dacă da, coboară. „Un singur focar, un traseu clar."',
            'reveal': 'Pe ecran: focarul atins primul și un traseu de citire deliberat.',
        },
        {
            'nume': 'ETAPA 4 · IERARHIA',
            'obiectiv': 'Facem ierarhia vizibilă instant cu mărime, poziție, contrast și spațiu, apoi grupăm ce ține împreună.',
            'stare': 'Rigoare. Importanța se vede instant, nu se caută.',
            'voce': '„Etapa 4: ierarhia. Greutate inegală pentru lucruri inegale. Două sau trei niveluri: focar, suport, detaliu. Și punem împreună ce e legat."',
            'demo': 'Pe ecran: mărim focarul, scădem detaliul, grupăm obiectele aceleiași întrebări, separăm ce nu e legat.',
            'ai': 'AI-ul propune niveluri și grupuri. Noi verificăm că ierarhia servește decizia, nu estetica.',
            'control': 'Confirmăm: trei niveluri clare, grupuri coerente. „Proximitatea spune adevărul despre relații: lângă înseamnă împreună."',
            'reveal': 'Pe ecran: o pagină cu ierarhie clară și grupuri care se citesc dintr-o privire.',
        },
        {
            'nume': 'ETAPA 5 · SPAȚIUL',
            'obiectiv': 'Folosim spațiul gol ca instrument, apoi aliniem toată pagina la răspunsul venit din analiză. Promptul 2.',
            'stare': 'Încredere. Compoziția nu e estetică; servește răspunsul.',
            'voce': '„Etapa 5: spațiul. Golul nu e lipsă, e instrument: izolează focarul. Apoi aliniem toată pagina la răspuns. Build-up guvernat, nu orb."',
            'demo': 'Pe ecran: lăsăm gol în jurul focarului, rulăm Promptul 2, corectăm varianta AI după se vede întâi ce decide.',
            'ai': 'Promptul 2 propune o variantă de compoziție. Noi o aliniem la răspuns: răspunsul decide ce e focar, ce e suport.',
            'control': 'Confirmăm: spațiul lucrează, compoziția servește răspunsul. „Aceleași obiecte, alta așezare, alta decizie."',
            'reveal': 'Pe ecran: pagina aliniată la răspuns, cu spațiul gol care întărește ierarhia.',
        },
        {
            'nume': 'ETAPA 6 · PROBA',
            'obiectiv': 'Trecem pagina prin testul celui de-al doilea ochi și o predăm pentru extragerea mesajului.',
            'stare': 'Claritate. Pagina nu mai e un zid; ochiul prinde întâi ce contează.',
            'voce': '„Etapa 6: proba. Dăm pagina cuiva trei secunde. Prinde aceeași prioritate? Dacă da, trece. Nu mai aranjăm; conducem privirea decidentului."',
            'demo': 'Pe ecran: un cititor proaspăt vede pagina trei secunde; predarea paginii coerente către treapta de sintetizare.',
            'ai': 'AI-ul rezumă compoziția. Noi confirmăm că am așezat pagina, nu am formulat mesajul.',
            'control': 'Confirmăm: focar prins, ordine corectă, pagină coerentă. „C14 așază pagina. Treapta de sintetizare îi extrage mesajul."',
            'reveal': 'Pe ecran: pagina coerentă, predată; aceleași grafice, acum o decizie dintr-o privire.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · PAGINA COMPUSĂ, PREDATĂ')
    txt('C14 așază spațial obiectele vizuale oneste primite de la treapta de vizualizare. Pagina '
        'are un focar, o ierarhie și un traseu de citire, testate cu al doilea ochi. C14 face '
        'pagina coerentă; extragerea și formularea mesajului încep la treapta de sintetizare. '
        'Predăm o pagină de raport, cu valorile sursă neatinse și suma conservată cap-coadă.')
    p = doc.add_paragraph()
    r = p.add_run('Ce vede ochiul întâi schimbă decizia.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))


if __name__ == '__main__':
    main()
