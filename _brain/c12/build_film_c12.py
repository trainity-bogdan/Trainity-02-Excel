#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c12.py - FILM-Excel-12-Interpretare.docx (script video procedural).

Structura = c09/c10 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C12 INTERPRETARE.
Garda C12: a explica (De ce?). Clasamentul = INPUT citit (C11), zero re-ierarhizare,
zero tablou vizual de raportare (T4), zero what-if/simulare/actiune (T5). Inchide T3.
Zero cifre business (R-V02.15).
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'c12/FILM-Excel-12-Interpretare.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C12 INTERPRETARE')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 3 ANALIZA · Construcția 12 din 20 · închide treapta T3')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK→DEMO→CONTROL→REVEAL.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('Știi care a crescut. Nu știi de ce. Astăzi transformăm rezultatul în explicație.')
    lbl('MIZA')
    txt('Trecem de la rezultatul numeric la explicația verbală: cauza citită din model, nu '
        'ghicită, ancorată în relații, măsuri și comparații, scrisă într-o frază pe care un '
        'om o poate înțelege, contesta și verifica. O explicație bună spune de ce, nu doar cât.')
    lbl('AHA (INSIGHT)')
    txt('Nu rezultatul contează. Contează de ce apare rezultatul.')
    lbl('MANTRA · TRAINITY')
    txt('Cifra spune cât. Explicația spune de ce.')
    lbl('WOW (PAYOFF FINAL)')
    txt('Același clasament putea avea două explicații opuse. Acum doar una se verifică în '
        'date, și o poți arăta oricui.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Nu citi rezultatul. Explică-l.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion (2-3 cuvinte). FRAZĂ DE IMPACT = '
        'exec-phrase. Sursă pentru slide-urile exec din HTML-Video.')
    execs = [
        ('1 · REALITATE', 'Rezultat fără cauză',
         'Te uiți la un clasament corect. Și la o întrebare fără răspuns: de ce?'),
        ('2 · INVESTIGAȚIE', 'Cauza, citită nu ghicită',
         'Nu ghicim de ce. Citim cauza din model și separăm coincidența de cauză.'),
        ('3 · TRANSFORMARE', 'Rezultatul devine explicație',
         'Coborâm sub total, găsim cauza și o scriem ca frază verificabilă.'),
        ('4 · VERIFICARE', 'Explicația rezistă',
         'O explicație pe care nu o poți arăta în date nu trece. A noastră se arată.'),
        ('5 · STABILIZARE', 'Explicație ancorată',
         'Cauza reală rămâne și pe datele de luna viitoare. Nu se agață de o poveste.'),
        ('6 · CONFIRMARE', 'Analiză închisă',
         'Erau cifre. Acum sunt explicate, citit din model. Treapta T3 e finalizată.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C12')
    txt('Construcția 12 INTERPRETARE preia clasamentul de la construcția precedentă. '
        'Comparațiile au arătat CARE conduce și cu cât diferă: avem un rezultat numeric corect. '
        'Acum nu mai întrebăm care, ci de ce: și transformăm rezultatul în explicație.')
    txt('Pentru prima dată, cifra nu mai e doar o constatare. E o explicație verbală, cu o '
        'cauză citită din model, pe care o poți arăta înapoi în date.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Citește din model cauzele posibile: ce felie, ce factor, ce tipar pare să tragă '
        'rezultatul, și ce apare doar împreună fără să fie cauză. Nu ghicește; citește și semnalează.')
    lbl('EXCEL · INTERPRETARE')
    txt('Găzduiește foaia Interpretare: cauza citită din model prin drill sub total, factorii '
        'principali, insight-ul verbal verificabil. Valorile sursă rămân neatinse.')
    lbl('CONTROL UMAN')
    txt('Formulăm întrebarea „de ce", separăm coincidența de cauză, respingem narativul '
        'nesusținut, scriem insight-ul verificabil. Nu re-clasăm și nu raportăm vizual: explicăm.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CELE 5 FENOMENE')
    fenomene = [
        ('01 · CLASAMENT FĂRĂ CAUZĂ',
         'Vezi care conduce și cu cât; fără cauza din spate, clasamentul arată CARE, nu DE CE.'),
        ('02 · ÎMPREUNĂ NU ÎNSEAMNĂ DIN CAUZA',
         'Două lucruri cresc în același timp; ce apare împreună nu explică automat ce produce rezultatul.'),
        ('03 · POVESTEA PLAUZIBILĂ',
         'O explicație sună perfect, dar nimeni nu o poate arăta în date: o presupunere, nu un insight.'),
        ('04 · CAUZA UNICĂ FORȚATĂ',
         'Se caută un singur vinovat unde lucrează mai mulți factori; explicația unică e comodă și greșită.'),
        ('05 · CAUZA ASCUNSĂ DE TOTAL',
         'Totalul arată direcția; cauza reală apare doar când cobori sub el, pe tăietura potrivită.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 5 fenomene interpretative. Rezultatul e pregătit să devină explicație verificabilă.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · REALITATE',
            'obiectiv': 'Confirmăm clasamentul predat de comparații, formulăm întrebarea „de ce?", vedem că un clasament arată CARE, nu DE CE.',
            'stare': 'Atenție. Un rezultat corect, dar fără cauză: clasamentul nu spune încă de ce.',
            'voce': '„Etapa 1: realitate. Avem un clasament corect. Întrebarea de business e: de ce? Răspunsul trebuie să fie o explicație, nu o altă cifră."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C12, foaia Comparatii și foaia Interpretare; arătăm diferența dintre o constatare și o explicație.',
            'ai': 'AI-ul arată ce rezultate cer un „de ce". Noi alegem rezultatul de explicat, nu o nouă cifră.',
            'control': 'Punem întrebarea-test: de ce conduce acela? „Un clasament fără cauză nu e o explicație. O citim."',
            'reveal': 'Pe ecran: clasamentul gata, cu o întrebare „de ce?" care așteaptă o cauză.',
        },
        {
            'nume': 'ETAPA 2 · INVESTIGAȚIE',
            'obiectiv': 'Promptul 1 citește cauza din model: ce felie trage rezultatul, ce apare doar împreună, pe ce tăietură cobori.',
            'stare': 'Atenție la sens. Înainte să explicăm, întrebăm modelul ce cauze susțin datele.',
            'voce': '„Etapa 2: investigație. Citim cauza din model: ce factor pare să tragă rezultatul și ce e doar coincidență."',
            'demo': 'Pe ecran: rulăm Promptul 1; AI-ul citește din date factorii posibili și semnalează coincidențele.',
            'ai': 'AI-ul citește cauzele posibile. Noi separăm ce apare împreună de ce produce rezultatul.',
            'control': 'Verificăm că factorul reținut chiar produce rezultatul. „Cauză citită. Mergem mai departe."',
            'reveal': 'Pe ecran: rezultatul cu factorii principali candidați și coincidențele puse deoparte.',
        },
        {
            'nume': 'ETAPA 3 · TRANSFORMARE',
            'obiectiv': 'Coborâm sub total, găsim cauza ascunsă de agregare, scriem insight-ul ca frază verificabilă. Promptul 2.',
            'stare': 'Construcție calmă. Rezultatul devine explicație: o frază, nu o părere.',
            'voce': '„Etapa 3: transformare. Coborâm sub total, găsim cauza reală și o scriem într-o frază pe care o poți arăta în date."',
            'demo': 'Pe ecran: în foaia Interpretare, drill sub categoria care conduce; insight-ul scris cu felia care îl susține.',
            'ai': 'Promptul 2 ne spune pe ce tăietură coborâm și formulează insight-ul. Noi separăm explicația de clasament.',
            'control': 'Confirmăm: o frază clară, cu o cauză și o felie care o dovedește. „Insight scris. Predăm la verificare."',
            'reveal': 'Pe ecran: rezultatul de la început, acum o explicație verbală ancorată în model.',
        },
        {
            'nume': 'ETAPA 4 · VERIFICARE',
            'obiectiv': 'Verificăm că explicația se arată în date, respingem narativul nesusținut, confirmăm că nu am forțat o cauză unică.',
            'stare': 'Rigoare. O explicație pe care nu o poți dovedi în date nu trece.',
            'voce': '„Etapa 4: verificare. Pentru fiecare insight arătăm felia care îl susține. Ce nu se confirmă, se respinge."',
            'demo': 'Pe ecran: fiecare frază de insight trimite la felia, măsura sau comparația din model care o probează.',
            'ai': 'AI-ul semnalează unde o explicație nu are sprijin în date. Noi respingem narativul plauzibil dar nesusținut.',
            'control': 'Reconciliem: fiecare „de ce" e dovedit, factorii principali numiți, nicio cauză unică forțată. „Explicație verificabilă."',
            'reveal': 'Pe ecran: insight-urile, fiecare cu dovada lui în model.',
        },
        {
            'nume': 'ETAPA 5 · STABILIZARE',
            'obiectiv': 'Ancorăm explicația la model. Un rând nou nu schimbă cauza dacă e reală. Regula anti-poveste.',
            'stare': 'Încredere. O explicație bună rămâne valabilă și pe datele noi.',
            'voce': '„Etapa 5: stabilizare. Insight-ul trăiește lângă felia care îl susține; un rând nou nu îl răstoarnă dacă e cauză reală."',
            'demo': 'Pe ecran: adăugăm date noi; cauza reală rămâne, o coincidență prinsă la moment ar fi căzut.',
            'ai': 'AI-ul confirmă că explicația ține pe datele noi. Noi notăm regula anti-poveste.',
            'control': 'Verificăm: cauza citită din model se confirmă sau se ajustează cinstit. „Explicație stabilă."',
            'reveal': 'Pe ecran: aceeași explicație care rezistă peste datele noi.',
        },
        {
            'nume': 'ETAPA 6 · CONFIRMARE',
            'obiectiv': 'Prima citire a insight-ului confirmă setul de explicații ancorate. Închidem analiza: treapta T3 finalizată.',
            'stare': 'Claritate. Rezultatele nu mai sunt simple constatări; sunt explicate.',
            'voce': '„Etapa 6: confirmare. Întrebi de ce, modelul explică. Aceeași întrebare, aceeași explicație. Analiza setului e completă."',
            'demo': 'Pe ecran: prima citire a insight-ului, cu cauza și felia care o susține; închiderea treptei T3.',
            'ai': 'AI-ul rezumă setul de explicații. Noi confirmăm că am explicat, nu am re-clasat.',
            'control': 'Confirmăm: fiecare insight are întrebare, cauză citită din model, dovadă. „Analiză închisă. Treapta T3 finalizată."',
            'reveal': 'Pe ecran: setul de explicații ancorate; setul a fost corect, înțeles și interpretat.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · TREAPTA T3 FINALIZATĂ')
    txt('C12 explică ce s-a întâmplat și de ce, pe baza modelului. Am completat analiza '
        'setului: treapta T3 este finalizată. Setul a fost legat, măsurat, comparat și acum '
        'interpretat: de la „ce există" la „de ce arată așa". Cum se vede dintr-o privire și '
        'cum se pune în pagină începe la treapta de raportare. Predăm un set de explicații '
        'ancorate, cu valorile sursă neatinse și suma conservată cap-coadă.')
    p = doc.add_paragraph()
    r = p.add_run('Nu citi rezultatul. Explică-l.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))


if __name__ == '__main__':
    main()
