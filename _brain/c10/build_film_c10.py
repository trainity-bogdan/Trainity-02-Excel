#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c10.py - FILM-Excel-10-Masuri.docx (script video procedural).

Structura = c09 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C10 MASURI.
Garda C10: a defini (Cat?). Zero ranking/contributie (C11), cauza (C12), dashboard (T4).
Zero cifre business (R-V02.15).
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'c10/FILM-Excel-10-Masuri.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C10 MĂSURI')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 3 ANALIZA · Construcția 10 din 20')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK→DEMO→CONTROL→REVEAL.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('Ai toate cifrele. Și nicio decizie. Astăzi le transformăm în măsuri definite.')
    lbl('MIZA')
    txt('Transformăm cifrele brute în măsuri utile, controlabile, explicabile: o singură '
        'definiție, o bază de raportare clară, un reper. O măsură bună reduce haosul, nu îl amplifică.')
    lbl('AHA (INSIGHT)')
    txt('Nu cifra contează. Contează ce întrebare răspunde cifra.')
    lbl('MANTRA · TRAINITY')
    txt('Măsura corectă răspunde întrebării corecte.')
    lbl('WOW (PAYOFF FINAL)')
    txt('Același set de date putea susține concluzii opuse. Acum cifra care primește '
        'autoritate este o măsură definită, și nu mai minte.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Nu calcula mai mult. Măsoară mai corect.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion (2-3 cuvinte). FRAZĂ DE IMPACT = '
        'exec-phrase. Sursă pentru slide-urile exec din HTML-Video.')
    execs = [
        ('1 · REALITATE', 'Cifre fără măsură',
         'Te uiți la un model care răspunde. Și la cifre care încă nu sunt măsuri.'),
        ('2 · INVESTIGAȚIE', 'Întrebarea înaintea cifrei',
         'Nu calculăm tot ce se poate. Întrebăm întâi ce măsură răspunde.'),
        ('3 · TRANSFORMARE', 'Măsura definită',
         'O definim o dată. Cu o bază declarată și un reper.'),
        ('4 · VERIFICARE', 'Aceeași pe orice felie',
         'O măsură care se rupe la filtru nu e măsură. A noastră ține.'),
        ('5 · STABILIZARE', 'Definiție ancorată',
         'O măsură bună se definește o dată la sursă și rămâne.'),
        ('6 · CONFIRMARE', 'Măsuri controlabile',
         'Erau cifre brute. Acum sunt măsuri cu bază și reper, predate mai departe.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C10')
    txt('Construcția 10 MĂSURI preia modelul de la construcția precedentă. C09 a legat '
        'tabelele prin chei: avem un model interogabil care răspunde la întrebări. Acum nu '
        'mai întrebăm dacă poate răspunde, ci cât: și transformăm răspunsul într-o măsură.')
    txt('Pentru prima dată, cifra nu mai e un calcul de moment. E o măsură definită o '
        'singură dată, cu bază și reper, aceeași pe orice felie.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Propune măsura potrivită întrebării: ce agregăm, care e baza de raportare, ce '
        'cifre nu merită promovate. Nu definește; descrie și recomandă.')
    lbl('EXCEL · MĂSURI')
    txt('Găzduiește definiția măsurii ca sursă unică de adevăr, peste modelul legat. '
        'Recalculează corect pe orice tăietură. Valorile sursă rămân neatinse.')
    lbl('CONTROL UMAN')
    txt('Formulăm întrebarea, alegem măsura, declarăm baza, ancorăm reperul, verificăm '
        'robustețea. Nu comparăm și nu clasăm: definim.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CELE 5 FENOMENE')
    fenomene = [
        ('01 · TOTAL FĂRĂ ÎNTREBARE',
         'Se adună o coloană fiindcă se poate; fără o întrebare clară, totalul nu răspunde la nimic.'),
        ('02 · MEDIA CARE ASCUNDE',
         'O medie pe tot setul amestecă lucruri diferite și poate descrie situații opuse.'),
        ('03 · PROCENT FĂRĂ BAZĂ',
         'Un procent fără numitor declarat pare precis și este gol: procent din ce?'),
        ('04 · TREI CIFRE PENTRU O ÎNTREBARE',
         'Aceeași măsură recalculată în trei locuri dă trei valori; lipsește sursa unică de adevăr.'),
        ('05 · MĂSURA FRAGILĂ',
         'Pui un filtru și măsura se rupe; o măsură care nu rezistă la context nu e o măsură.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 5 fenomene cantitative. Cifrele sunt pregătite să devină măsuri definite.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · REALITATE',
            'obiectiv': 'Confirmăm modelul legat de C09, formulăm întrebarea „cât?", vedem că o cifră brută nu e încă o măsură.',
            'stare': 'Atenție. Un model care răspunde, dar cifrele lui sunt încă brute, fără definiție.',
            'voce': '„Etapa 1: realitate. Avem un model legat corect. Întrebarea de business e simplă: cât? Răspunsul trebuie să fie o măsură, nu o cifră de moment."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C10, foaia Vanzari și foaia Masuri; arătăm diferența dintre o cifră calculată ad-hoc și o măsură definită.',
            'ai': 'AI-ul arată ce întrebări de business pot fi măsurate. Noi alegem întrebarea, nu calculul.',
            'control': 'Punem întrebarea-test: cât în medie? „O cifră fără definiție nu e o măsură. O definim."',
            'reveal': 'Pe ecran: modelul gata, cu o întrebare „cât?" care așteaptă o măsură.',
        },
        {
            'nume': 'ETAPA 2 · INVESTIGAȚIE',
            'obiectiv': 'Promptul 1 identifică măsura potrivită: ce agregăm, ce bază de raportare, ce cifre rămân decorative.',
            'stare': 'Atenție la sens. Înainte să calculăm, întrebăm modelul ce măsură răspunde corect.',
            'voce': '„Etapa 2: investigație. Întrebăm ce măsură răspunde la întrebare: total, medie sau rată, și care e baza de raportare."',
            'demo': 'Pe ecran: rulăm Promptul 1; AI-ul propune măsura, numitorul și cifrele care nu merită promovate.',
            'ai': 'AI-ul propune măsura și baza. Noi distingem cifra utilă de cifra decorativă.',
            'control': 'Verificăm că măsura propusă răspunde la o întrebare reală. „Măsură aleasă. Mergem mai departe."',
            'reveal': 'Pe ecran: întrebarea de business cu măsura candidat și baza ei de raportare.',
        },
        {
            'nume': 'ETAPA 3 · TRANSFORMARE',
            'obiectiv': 'Definim măsura o singură dată, declarăm baza de raportare, ancorăm într-un reper. Promptul 2.',
            'stare': 'Construcție calmă. Cifra brută devine măsură: o definiție, nu un calcul repetat.',
            'voce': '„Etapa 3: transformare. Definim măsura o singură dată, ca sursă unică de adevăr, cu baza ei și un reper."',
            'demo': 'Pe ecran: scriem definiția în foaia Masuri; o medie devine total împărțit la o bază declarată; fixăm un reper.',
            'ai': 'Promptul 2 ne spune definiția, baza, reperul și semnalul. Noi separăm definirea de comparație.',
            'control': 'Confirmăm: o singură definiție, o bază clară, un reper. „Măsură definită. Predăm la verificare."',
            'reveal': 'Pe ecran: cifra brută de la început, acum o măsură cu definiție, bază și reper.',
        },
        {
            'nume': 'ETAPA 4 · VERIFICARE',
            'obiectiv': 'Verificăm robustețea: aceeași definiție pe orice tăietură, baza declarată, semnal explicabil față de reper.',
            'stare': 'Rigoare. O măsură care se rupe la primul filtru nu e o măsură.',
            'voce': '„Etapa 4: verificare. Recalculăm aceeași măsură pe categorie, client, perioadă. Valorile diferă fiindcă felia diferă, nu măsura."',
            'demo': 'Pe ecran: în foaia Masuri_Context, aceeași definiție dă rezultat corect pe patru tăieturi diferite.',
            'ai': 'AI-ul semnalează unde o măsură și-ar pierde baza. Noi confirmăm numitorul și reperul.',
            'control': 'Reconciliem: aceeași definiție peste tot, semnal explicabil față de reper. „Măsură robustă."',
            'reveal': 'Pe ecran: măsura citită pe mai multe felii, identică în definiție, corectă în rezultat.',
        },
        {
            'nume': 'ETAPA 5 · STABILIZARE',
            'obiectiv': 'Ancorăm definiția. Un rând nou se include singur în măsură. Regula anti-derivă.',
            'stare': 'Încredere. O măsură bună rămâne valabilă când apar rânduri sau surse noi.',
            'voce': '„Etapa 5: stabilizare. Definiția trăiește într-un singur loc; un rând nou se recalculează singur prin ea."',
            'demo': 'Pe ecran: adăugăm o tranzacție nouă; măsura se recalculează singură, fără să rescriem definiția.',
            'ai': 'AI-ul confirmă că măsura a absorbit rândul nou. Noi notăm regula anti-derivă.',
            'control': 'Verificăm: definiția, baza și reperul rămân valabile la sursa nouă. „Măsură stabilă."',
            'reveal': 'Pe ecran: aceeași definiție care lucrează singură peste datele noi.',
        },
        {
            'nume': 'ETAPA 6 · CONFIRMARE',
            'obiectiv': 'Prima citire a măsurii confirmă setul. Valorile sursă neatinse. Predăm către C11.',
            'stare': 'Claritate. Cifrele nu mai sunt brute; sunt măsuri controlabile.',
            'voce': '„Etapa 6: confirmare. Cerem măsura și primim o valoare, baza ei și semnalul. Aceeași întrebare, același răspuns."',
            'demo': 'Pe ecran: prima citire a măsurii; răspunsul vine cu baza și semnalul față de reper.',
            'ai': 'AI-ul rezumă setul de măsuri. Noi confirmăm că am definit, nu am comparat.',
            'control': 'Confirmăm: fiecare măsură are întrebare, definiție, bază, reper. „Set de măsuri controlabile."',
            'reveal': 'Pe ecran: setul de măsuri definite, gata de predat către C11 pentru comparație.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · PREDARE CĂTRE C11')
    txt('C10 definește măsuri. Atât. Care măsură e mai mare, care contribuie mai mult, ce '
        'loc ocupă fiecare felie, începe la C11. Predăm un set de măsuri definite, cu '
        'valorile sursă neatinse și suma conservată cap-coadă.')
    p = doc.add_paragraph()
    r = p.add_run('Nu calcula mai mult. Măsoară mai corect.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))


if __name__ == '__main__':
    main()
