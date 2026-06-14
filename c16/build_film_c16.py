#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c16.py - FILM-Excel-16-Livrare.docx (script video procedural).

Structura = c14 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C16 LIVRARE.
Garda C16: a livra raportul ca obiect de decizie. ZERO sinteza mesaj (C15),
ZERO layout/pagina (C14), ZERO sistem recurent (C17), ZERO logistica.
Cifre in Excel (R-V02.15).
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'c16/FILM-Excel-16-Livrare.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C16 LIVRARE')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 4 RAPORTARE · Construcția 16 din 20 · închide treapta T4')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK->DEMO->CONTROL->REVEAL.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('Raportul e gata. Decizia, nu. Astăzi dăm raportului forma care îl face o decizie gata de luat, nu doar informație trimisă.')
    lbl('MIZA')
    txt('Valoarea întregului lanț se realizează sau se pierde în momentul livrării. Trecem de la '
        'un raport complet, dar care nu mișcă, la o foaie-raport de decizie: concluzia cerută în cap, '
        'cadrul de decizie explicit, predabilă fără autor lângă, citită la fel de oricine.')
    lbl('AHA (INSIGHT)')
    txt('Un raport care nu produce o decizie nu e livrat, e doar trimis.')
    lbl('MANTRA · TRAINITY')
    txt('Nu livrez informație. Livrez o decizie gata de luat.')
    lbl('WOW (PAYOFF FINAL)')
    txt('Același raport: trimis, stătea pe loc; livrat ca decizie, se hotărăște dintr-o citire.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Raportul care decide.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion. FRAZĂ DE IMPACT = exec-phrase.')
    execs = [
        ('1 · REALITATE', 'Raportul care nu mișcă',
         'Te uiți la un raport complet. Și la o întrebare fără răspuns: de ce nu se decide din el?'),
        ('2 · INVESTIGAȚIE', 'Ce cere decizia',
         'Nu tot raportul susține decizia. O parte e semnal, restul e zgomot care aglomerează.'),
        ('3 · TRANSFORMARE', 'Forma de decizie',
         'Concluzia urcă în cap, cadrul decizie, risc, concluzie, pas se scrie, detaliul coboară în anexă.'),
        ('4 · VERIFICARE', 'Foaia se susține singură',
         'Decidentul hotărăște doar din foaie. Dacă te cheamă pentru context de bază, încă e informație.'),
        ('5 · STABILIZARE', 'Regula, nu inspirația',
         'Cele șase reguli fac orice raport următor decision-ready, complet și fără surplus.'),
        ('6 · CONFIRMARE', 'Decizia predată',
         'Forma de decizie scurtează drumul spre hotărâre. Predai o decizie, nu un teanc de date.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C16')
    txt('Construcția 16 LIVRARE preia mesajul de la sinteză. Treptele precedente au dat un '
        'rezultat corect, l-au vizualizat, l-au compus și au sintetizat mesajul esențial: îl avem, '
        'dar raportul e mut pentru decizie. Un decident citește totul și tot nu hotărăște. Acum nu '
        'mai sintetizăm și nu re-formulăm; îi dăm raportului o formă de decizie.')
    txt('Pentru ultima dată în treapta de raportare, forma nu mai e un detaliu. E o decizie: '
        'același mesaj, livrat ca decizie, se hotărăște dintr-o citire. C16 face raportul o decizie de luat.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Propune ce decizie se cere din raport și ce e relevant pentru ea, apoi structurează '
        'foaia. Nu hotărăște și nu re-sintetizează mesajul; aceea rămâne la om.')
    lbl('EXCEL · LIVRARE')
    txt('Găzduiește foaia Livrare: decizia cerută, cadrul de decizie peste agregatul real, ce '
        'intră în decizie vs anexă, verificarea self-standing, cele șase reguli. Valorile sursă rămân neatinse.')
    lbl('CONTROL UMAN')
    txt('Numim decizia, scriem cadrul de decizie, punem pragul de risc și pasul următor, verificăm '
        'foaia fără autor. Nu sintetizăm mesajul și nu compunem pagina: dăm forma de decizie.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CELE 6 FENOMENE')
    fenomene = [
        ('01 · CONCLUZIA ÎNGROPATĂ',
         'Concluzia relevantă stă în tabel, la pagina paisprezece, nu în capul foii. Regula: concluzia cerută urcă în prima linie, ca decizie.'),
        ('02 · DETALIUL CARE AGLOMEREAZĂ',
         'Tot detaliul brut, linie cu linie, sufocă decizia în context. Regula: detaliul coboară în anexă, sus rămâne agregatul de decizie.'),
        ('03 · LIPSA CADRULUI',
         'Doar cifre, fără decizie, risc, concluzie și pas următor: e informație, nu decizie. Regula: se adaugă cadrul de decizie, explicit.'),
        ('04 · RISCUL IMPLICIT',
         'Riscul e lăsat pe seama decidentului, să-l deducă din cifre. Regula: riscul se scrie, cu pragul lui declarat.'),
        ('05 · OPRIT LA CONSTATARE',
         'Raportul se termină în asta e situația, fără nicio acțiune propusă. Regula: se scrie pasul următor, ancorat în date.'),
        ('06 · NEVOIE DE AUTOR',
         'Fără unitate, perioadă și referință, raportul cere autorul lângă el. Regula: fiecare cifră poartă unitate, perioadă și referință.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 6 fenomene de livrare, fiecare cu corecția lui decision-ready. Raportul e pregătit '
        'să devină o foaie-raport de decizie.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · REALITATE',
            'obiectiv': 'Preluăm mesajul sintetizat și rezultatul corect, constatăm că raportul e complet dar nu produce o decizie, ne pregătim să-i dăm forma de decizie (nu să-l re-sintetizăm).',
            'stare': 'Atenție. Un raport complet, dar care nu mișcă: are tot, dar nu are decizia.',
            'voce': '„Etapa 1: realitate. Avem mesajul sintetizat și răspunsul corect. Dar raportul e mut pentru decizie: decidentul tot întreabă. Îi dăm forma de decizie."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C16, foaia Livrare; arătăm raportul complet care încă nu produce o hotărâre.',
            'ai': 'AI-ul arată că mesajul există deja, sintetizat. Noi alegem decizia de făcut posibilă, fără să re-sintetizăm.',
            'control': 'Punem regula-test: dăm formă de decizie, nu re-sintetizăm. „Un raport complet, dar care nu produce o decizie, nu e încă livrat."',
            'reveal': 'Pe ecran: raportul complet, încă fără cadru de decizie, gata de transformat.',
        },
        {
            'nume': 'ETAPA 2 · INVESTIGAȚIE',
            'obiectiv': 'Descoperim ce decizie se cere din raport și ce e semnal vs zgomot pentru decizie. Promptul 1.',
            'stare': 'Atenție la decizie. Înainte să construim foaia, vedem ce-i trebuie decidentului ca să hotărască.',
            'voce': '„Etapa 2: investigație. Nu tot raportul susține decizia. O parte e semnal, restul e zgomot. Decizia cerută o numim noi."',
            'demo': 'Pe ecran: rulăm Promptul 1; AI-ul propune ce decizie se cere, ce risc o însoțește, ce e relevant vs zgomot.',
            'ai': 'AI-ul propune lista. Noi judecăm ce intră în decizie; nu lăsăm AI-ul să hotărască.',
            'control': 'Separăm semnalul de zgomot pe viu: „Aceleași date, dar o parte susține hotărârea, restul doar aglomerează."',
            'reveal': 'Pe ecran: raportul împărțit în ce intră în decizie și ce rămâne anexă.',
        },
        {
            'nume': 'ETAPA 3 · TRANSFORMARE',
            'obiectiv': 'Construim cadrul de decizie: concluzia urcă în cap, scriem decizie, risc, concluzie, pas următor, detaliul coboară în anexă. Promptul 2.',
            'stare': 'Construcție calmă. Raportul capătă forma de decizie.',
            'voce': '„Etapa 3: transformare. Concluzia urcă în prima linie, scriem cadrul de decizie, detaliul coboară în anexă. Nu re-formulăm mesajul, îl poziționăm."',
            'demo': 'Pe ecran: rulăm Promptul 2; AI-ul propune structura foii, noi punem pragul de risc și pasul următor.',
            'ai': 'Promptul 2 așază cadrul. Noi punem decizia de luat, riscul cu pragul, pasul următor ca acțiune.',
            'control': 'Confirmăm: o singură foaie-raport de decizie, nu un dashboard. „Raportul a trecut de la informație completă la decizie de luat."',
            'reveal': 'Pe ecran: raportul de la început, acum o foaie-raport de decizie.',
        },
        {
            'nume': 'ETAPA 4 · VERIFICARE',
            'obiectiv': 'Testăm foaia ca obiect de decizie: poate decidentul hotărî doar din ea, fără autor; riscul și pasul scrise explicit.',
            'stare': 'Rigoare. O foaie decision-ready se decide fără autorul lângă ea.',
            'voce': '„Etapa 4: verificare. Punem foaia în fața decidentului și plecăm. Dacă ne cheamă pentru context de bază, încă e informație."',
            'demo': 'Pe ecran: dăm foaia cuiva care vede doar raportul; ascultăm ce întrebări apar.',
            'ai': 'AI-ul semnalează unde decidentul ar trebui să deducă; noi scriem explicit riscul și pasul.',
            'control': 'Reconciliem: întrebările de bază eliminate, rămân doar cele inevitabile de decizie. „Foaie verificată, nu doar completă."',
            'reveal': 'Pe ecran: foaia din care decidentul hotărăște dintr-o citire.',
        },
        {
            'nume': 'ETAPA 5 · STABILIZARE',
            'obiectiv': 'Fixăm cele șase reguli ale foii-raport de decizie ca standard, tăiem surplusul, obținem o foaie repetabilă.',
            'stare': 'Încredere. Livrarea decision-ready nu e un noroc de o dată; e o regulă.',
            'voce': '„Etapa 5: stabilizare. Fixăm cele șase reguli. O foaie următoare, din aceeași natură, e decision-ready din naștere, nu din noroc."',
            'demo': 'Pe ecran: aplicăm aceleași reguli pe un raport nou; tăiem ce nu intră în decizie.',
            'ai': 'AI-ul aplică regulile pe un caz nou. Noi confirmăm că foaia se susține singură, fără surplus.',
            'control': 'Verificăm: cele șase reguli respectate, complet fără surplus. „Foaie-raport de decizie, repetabilă."',
            'reveal': 'Pe ecran: același standard de decizie, repetabil pe rapoarte noi.',
        },
        {
            'nume': 'ETAPA 6 · CONFIRMARE',
            'obiectiv': 'Forma de decizie scurtează drumul spre hotărâre, devenim cei care predau o decizie, predăm raportul către autonomie.',
            'stare': 'Claritate. Raportul nu mai e doar informație; produce o decizie.',
            'voce': '„Etapa 6: confirmare. Forma de decizie a scurtat drumul. Nu mai trimitem un raport; livrăm o decizie gata de luat."',
            'demo': 'Pe ecran: raportul-informație lângă raportul-decizie; predarea către treapta de autonomie.',
            'ai': 'AI-ul rezumă foaia-raport de decizie. Noi confirmăm că am livrat o decizie, nu am sistematizat.',
            'control': 'Confirmăm: decizie, risc și pas prezente, self-standing, gata. „C16 face raportul o decizie. Treapta de autonomie îl face să ruleze singur."',
            'reveal': 'Pe ecran: raportul-decizie predat; mesajul corect e acum și o hotărâre.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · DECIZIA PREDATĂ')
    txt('C16 dă raportului forma de decizie. Raportul-decizie e self-standing, complet fără '
        'surplus, predabil ca obiect de decizie. C16 face raportul o decizie; sistematizarea '
        'recurentă, automatizarea și guvernarea încep la treapta de autonomie. Predăm o foaie-raport '
        'de decizie, cu valorile sursă neatinse și suma conservată cap-coadă.')
    p = doc.add_paragraph()
    r = p.add_run('Raportul care decide.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))


if __name__ == '__main__':
    main()
