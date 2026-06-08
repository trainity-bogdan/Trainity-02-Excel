#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c13.py - FILM-Excel-13-Vizualizare.docx (script video procedural).

Structura = c12 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C13 VIZUALIZARE.
Garda C13: a da forma onesta. Obiect vizual onest, ZERO dashboard/pagina (C14),
zero mesaj (C15), zero livrare (C16), zero 'de ce' (C12). Cifre in Excel (R-V02.15).
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'c13/FILM-Excel-13-Vizualizare.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C13 VIZUALIZARE')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 4 RAPORTARE · Construcția 13 din 20 · deschide treapta T4')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK->DEMO->CONTROL->REVEAL.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('Cifra e corectă. Graficul minte. Astăzi dăm rezultatului o formă care îl arată adevărat.')
    lbl('MIZA')
    txt('O decizie e luată cu încredere pe o imagine falsă construită peste date corecte. '
        'Trecem de la rezultatul corect, dar mut în model, la un obiect vizual onest: forma '
        'aleasă, nu nimerită, verificată contra cifrei brute, citită la fel de oricine.')
    lbl('AHA (INSIGHT)')
    txt('Forma greșită minte cu cifra corectă.')
    lbl('MANTRA · TRAINITY')
    txt('Nu desenez frumos. Desenez adevărat.')
    lbl('WOW (PAYOFF FINAL)')
    txt('Aceeași cifră, două grafice, două concluzii opuse. Apoi forma onestă repară percepția.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Forma nu se nimerește. Se alege.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion. FRAZĂ DE IMPACT = exec-phrase.')
    execs = [
        ('1 · REALITATE', 'Răspuns fără formă',
         'Te uiți la un rezultat corect. Și la o întrebare fără răspuns: cum îl vede altcineva?'),
        ('2 · INVESTIGAȚIE', 'Forma e o alegere',
         'Aceeași cifră are mai multe forme. Una spune adevărul, alta minte cu cifra corectă.'),
        ('3 · TRANSFORMARE', 'Forma devine onestă',
         'Tipul urmează natura rezultatului, axa pleacă de la zero, codarea care minte iese.'),
        ('4 · VERIFICARE', 'Forma rezistă',
         'Vizualul citit singur dă aceeași concluzie ca cifra brută. Dacă spune mai mult, iese.'),
        ('5 · STABILIZARE', 'Forma devine regulă',
         'Cele șase reguli fac orice vizual următor onest din naștere, nu din noroc.'),
        ('6 · CONFIRMARE', 'Obiect onest',
         'Forma adevărată a reparat percepția. Obiectul vizual onest e gata de predat la compunere.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C13')
    txt('Construcția 13 VIZUALIZARE preia răspunsul de la analiză. Treapta precedentă a dat un '
        'rezultat corect, citit din model: îl avem, dar e mut. Un decident nu hotărăște privind '
        'un model. Acum nu mai producem niciun răspuns nou; îi dăm răspunsului o formă.')
    txt('Pentru prima dată, forma nu mai e un detaliu de stil. E o decizie: aceeași cifră, '
        'desenată altfel, spune altceva. C13 alege forma care o arată adevărată.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Propune forma care se potrivește naturii rezultatului și avertizează ce forme l-ar '
        'deforma. Generează vizualul repede. Nu garantează onestitatea lui; aceea o pune omul.')
    lbl('EXCEL · VIZUALIZARE')
    txt('Găzduiește foaia Vizualizare: rezultatul de vizualizat, forma onestă vs forma care '
        'minte, verificarea contra cifrei brute, cele șase reguli. Valorile sursă rămân neatinse.')
    lbl('CONTROL UMAN')
    txt('Alegem tipul după natura rezultatului, corectăm axa și scala, scoatem codarea care '
        'sugerează, verificăm vizualul contra cifrei. Nu compunem pagina și nu livrăm: dăm formă onestă.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CELE 6 FENOMENE')
    fenomene = [
        ('01 · AXA CARE EXAGEREAZĂ',
         'Axa nu pleacă de la zero și o diferență minusculă pare o prăpastie. Regula: axă de la zero sau abatere declarată.'),
        ('02 · TIPUL NEPOTRIVIT',
         'O evoluție pusă în plăcintă, categorii puse în linie: forma contrazice natura datelor. Regula: tipul urmează natura.'),
        ('03 · SCALA CARE INVENTEAZĂ',
         'O a doua axă sau o scală nedeclarată creează o corelație care nu există. Regula: o singură scală liniară, declarată.'),
        ('04 · CODAREA CARE SUGEREAZĂ',
         '3D, gradient sau arie dublă fac ochiul să citească o tendință inexistentă. Regula: o singură dimensiune codată coerent.'),
        ('05 · AGREGAREA CARE ASCUNDE',
         'Media netezește tot și ascunde variația care era mesajul. Regula: arăți distribuția, nu doar media.'),
        ('06 · ETICHETA LIPSĂ',
         'Fără unitate, etichetă sau context, cititorul ghicește ce vede. Regula: fiecare vizual poartă unitate, etichetă și context.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 6 fenomene de formă, fiecare cu regula lui de onestitate. Rezultatul e pregătit '
        'să capete un obiect vizual onest.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · REALITATE',
            'obiectiv': 'Preluăm răspunsul corect de la analiză, constatăm că e mut în model, ne pregătim să-i dăm formă (nu să-l re-analizăm).',
            'stare': 'Atenție. Un rezultat corect, dar invizibil: trăiește în model, nu în fața nimănui.',
            'voce': '„Etapa 1: realitate. Avem un răspuns corect. Dar e în model, mut. Un decident nu hotărăște dintr-un tabel. Îi dăm o formă."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C13, foaia Vizualizare; arătăm rezultatul corect care încă nu se vede.',
            'ai': 'AI-ul arată că răspunsul există deja. Noi alegem rezultatul de făcut vizibil, fără să inventăm altul.',
            'control': 'Punem regula-test: dăm formă, nu re-analizăm. „Un răspuns corect, dar invizibil, nu e încă o decizie. Îl facem văzut."',
            'reveal': 'Pe ecran: rezultatul corect, încă fără formă, gata de vizualizat.',
        },
        {
            'nume': 'ETAPA 2 · INVESTIGAȚIE',
            'obiectiv': 'Descoperim că aceeași cifră are mai multe forme și că forma greșită schimbă concluzia. Promptul 1.',
            'stare': 'Atenție la formă. Înainte să desenăm, vedem că forma e o alegere, nu o proprietate a cifrei.',
            'voce': '„Etapa 2: investigație. Aceeași cifră intră în multe forme. Una spune adevărul, alta minte cu cifra corectă. Forma o alegem noi."',
            'demo': 'Pe ecran: rulăm Promptul 1; AI-ul propune forma potrivită naturii rezultatului și ce forme l-ar deforma.',
            'ai': 'AI-ul propune forma. Noi judecăm dacă spune adevărul; desenăm aceeași cifră cu axa de la zero și cu axa trunchiată.',
            'control': 'Vedem minciuna de formă pe viu: două grafice, două concluzii. „Nu cifra s-a schimbat, ci forma. Și ochiul a crezut forma."',
            'reveal': 'Pe ecran: aceeași cifră, două forme, două concluzii opuse.',
        },
        {
            'nume': 'ETAPA 3 · TRANSFORMARE',
            'obiectiv': 'Dăm forma onestă: tipul urmează natura, corectăm axa și scala, scoatem codarea care sugerează. Promptul 2.',
            'stare': 'Construcție calmă. Rezultatul capătă forma care îl arată adevărat.',
            'voce': '„Etapa 3: transformare. Tipul urmează natura rezultatului, axa pleacă de la zero, codarea care minte iese. Forma nu se nimerește, se alege."',
            'demo': 'Pe ecran: rulăm Promptul 2; AI-ul generează vizualul, noi corectăm axa, scala, unitatea și scoatem 3D-ul.',
            'ai': 'Promptul 2 generează vizualul. Noi punem onestitatea: o singură scală, o singură dimensiune codată.',
            'control': 'Confirmăm: un singur obiect vizual onest, nu o pagină. „Forma a trecut de la arată bine la arată adevărat."',
            'reveal': 'Pe ecran: rezultatul de la început, acum un obiect vizual onest.',
        },
        {
            'nume': 'ETAPA 4 · VERIFICARE',
            'obiectiv': 'Verificăm vizualul contra cifrei brute, aplicăm testul celui de-al doilea ochi, marcăm forma care spune mai mult.',
            'stare': 'Rigoare. O formă onestă produce aceeași concluzie ca cifra brută.',
            'voce': '„Etapa 4: verificare. Punem vizualul lângă cifră. Dacă graficul spune mai mult sau altceva, forma minte și o corectăm."',
            'demo': 'Pe ecran: citim concluzia din grafic, apoi din număr; dăm vizualul cuiva care vede doar graficul.',
            'ai': 'AI-ul semnalează unde forma adaugă o concluzie nesusținută. Noi marcăm și corectăm.',
            'control': 'Reconciliem: vizual citit singur = cifra brută, cele șase întrebări trecute. „Formă verificată, nu doar frumoasă."',
            'reveal': 'Pe ecran: vizualul care dă aceeași concluzie ca cifra din care vine.',
        },
        {
            'nume': 'ETAPA 5 · STABILIZARE',
            'obiectiv': 'Fixăm cele șase reguli ca standard, punem eticheta, unitatea, contextul, obținem un obiect vizual repetabil.',
            'stare': 'Încredere. Onestitatea nu e un noroc de o dată; e o regulă.',
            'voce': '„Etapa 5: stabilizare. Fixăm cele șase reguli. Un vizual următor, din aceeași natură, se naște onest, nu din noroc."',
            'demo': 'Pe ecran: aplicăm aceleași reguli pe un rezultat nou; punem unitatea, eticheta, contextul.',
            'ai': 'AI-ul aplică regulile pe un caz nou. Noi confirmăm că obiectul se citește singur, nimic de ghicit.',
            'control': 'Verificăm: cele șase reguli respectate, etichete complete. „Obiect vizual onest, repetabil."',
            'reveal': 'Pe ecran: același standard de onestitate, repetabil pe rezultate noi.',
        },
        {
            'nume': 'ETAPA 6 · CONFIRMARE',
            'obiectiv': 'Forma onestă repară percepția, devenim garantul a ce vede altcineva, predăm obiectul către compunere.',
            'stare': 'Claritate. Rezultatul nu mai e mut; are o formă pe care ochiul o crede pentru că e adevărată.',
            'voce': '„Etapa 6: confirmare. Forma adevărată a reparat percepția. Nu mai facem grafice; răspundem de ce vede altcineva."',
            'demo': 'Pe ecran: forma care mințea lângă forma onestă; predarea obiectului vizual către treapta de compunere.',
            'ai': 'AI-ul rezumă obiectul vizual onest. Noi confirmăm că am dat formă, nu am compus pagina.',
            'control': 'Confirmăm: obiect verificat, etichetat, gata. „C13 face obiectul adevărat. Treapta de compunere îl așază în pagină."',
            'reveal': 'Pe ecran: obiectul vizual onest, predat; răspunsul corect e acum și vizibil.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · OBIECTUL ADEVĂRAT, PREDAT')
    txt('C13 dă forma onestă unui rezultat deja produs de analiză. Obiectul vizual e verificat '
        'contra cifrei brute, repetabil și complet etichetat. C13 face obiectul adevărat; '
        'compunerea paginii, ierarhia și ce vede ochiul întâi încep la treapta de compunere. '
        'Predăm un obiect vizual onest, cu valorile sursă neatinse și suma conservată cap-coadă.')
    p = doc.add_paragraph()
    r = p.add_run('Forma nu se nimerește. Se alege.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))


if __name__ == '__main__':
    main()
