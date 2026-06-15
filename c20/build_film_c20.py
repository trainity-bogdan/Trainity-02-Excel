#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_film_c20.py - FILM-Excel-20-Delegare.docx (script video procedural).

Structura = c19 FILM (HOOK->DEMO->CONTROL->REVEAL pe 6 etape), continut C20 DELEGARE.
Axa PREDARE / PROPRIETATE: rol care preia, harta de predare, comutatorul AUTOR_ACTIV,
parametru author-only vs PARAM_ documentat, V1-V4, STATUS NEPREDAT/PARTIAL/DELEGAT/AUTONOM,
drama V1 FAIL -> reparare -> AUTONOM. Verb DELEG. Artefact _DELEGARE (test VIU).
Garda C20: a preda proprietatea unui ROL, dovedit in workbook. ZERO management/HR/RACI/SOP,
ZERO delegare-ca-impartire-de-sarcini, ZERO control/reguli ca identitate (=C19, doar callback),
ZERO motor/automatizare ca identitate (=C18, doar callback). Rol, nu persoana.
Granita vs C19: C19 prinde un INPUT gresit in date; C20 prinde DEPENDENTA de autor.
C20 inchide PACK-ul C01-C20. Cifre in Excel (R-V02.15). Zero em-dash, zero en-dash.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'c20/FILM-Excel-20-Delegare.docx'

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def main():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    def h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        return p

    def h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        return p

    def lbl(t):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
        return p

    def txt(t):
        doc.add_paragraph(t)

    # ---- TITLU ----
    p = h1('FILM · C20 DELEGAREA SISTEMULUI DE LUCRU')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    txt('Pack 02 Excel · Treapta 5 AUTONOMIE · Construcția 20 din 20 (construcția-semnătură, închide pachetul)')
    txt('Script video procedural cinematic. 6 etape cu ritm HOOK->DEMO->CONTROL->REVEAL în interior.')
    txt('Arc: sistem încă al tău -> ce te face indispensabil -> harta de predare -> scoți autorul -> dependența ascunsă și repararea -> STATUS AUTONOM.')

    # ---- IDENTITATE CINEMATICA ----
    h1('IDENTITATE CINEMATICĂ')
    lbl('INTRIGA (HOOK)')
    txt('Cum deleg sistemul, ca să meargă fără mine? Sistemul tău rulează singur și se '
        'păzește singur. Și totuși e încă al tău, și numai al tău: la orice problemă, tot pe '
        'tine te sună. Un sistem care depinde de o singură persoană nu e autonom, e doar bine '
        'întreținut de acea persoană. Astăzi nu mai întrebăm „merge?". Predăm proprietatea.')
    lbl('BOMBĂ')
    txt('Sistemul tău merge fără tine. Dar la orice problemă, tot pe tine te sună.')
    lbl('MIZA')
    txt('Cât timp sistemul rămâne proprietatea informală a autorului, autonomia construită până '
        'aici e falsă: el merge doar cât autorul e disponibil să răspundă. Un sistem care rulează '
        'perfect, dar pe care îl deține o singură persoană, moare în ziua în care acea persoană '
        'pleacă, nu pentru că s-a stricat, ci pentru că nimeni altcineva nu îl deține, doar îl '
        'folosește. Delegarea mută proprietatea de la o persoană la un rol, și o face verificabil '
        'în workbook: nu „ți-am explicat, descurcă-te", ci „sistemul dovedește singur că rolul îl '
        'poate continua fără mine".')
    lbl('GREȘEALA CLASICĂ')
    txt('Crezi că ai delegat când ai explicat. Dar explicația nu mută proprietatea: la prima '
        'problemă, tot pe tine te sună.')
    lbl('AHA (INSIGHT)')
    txt('Un sistem nu e autonom pentru că merge singur. E autonom când îl poate deține altcineva.')
    lbl('MANTRA · TRAINITY')
    txt('Nu împart sarcini. Deleg sistemul.')
    lbl('WOW (PAYOFF FINAL)')
    txt('Apeși „scoate autorul", și sistemul nu se rupe: workbook-ul confirmă singur că alt rol '
        'îl poate continua.')
    lbl('MOTTO (SEMNĂTURĂ)')
    txt('Pleci, și munca nu mai e a ta.')

    # ---- SLIDES EXECUTIVE ----
    h1('SLIDES EXECUTIVE · SHOW FINAL VIDEO')
    txt('Cele 6 slide-uri din show-ul executiv de la finalul videoului (recapitulare '
        'cinematică, una per etapă). STARE = exec-emotion. FRAZĂ DE IMPACT = exec-phrase.')
    execs = [
        ('1 · REALITATE', 'Sistemul merge, dar e încă al tău',
         'Sistemul rulează și se păzește singur. Și totuși, la orice problemă, tot pe mine mă sună.'),
        ('2 · DEPENDENȚĂ', 'Ce mă face indispensabil',
         'Inventariez ce inputuri atârnă numai de mine, ținute în cap. Le marchez author-only.'),
        ('3 · HARTA', 'Harta de predare ca input al testului',
         'Rol, responsabilitate, acces pe zone, limite blocate, escaladare către un rol. Nu un tabel de citit.'),
        ('4 · TESTUL', 'Scoți autorul',
         'Apăs «scoate autorul». V1-V4 și STATUS se reașază singure. Nu bifez o predare, o dovedesc.'),
        ('5 · REPARARE', 'Dependența ascunsă, apoi repararea',
         'V1 aprinde FAIL: eu eram cheia. Mut parametrul pe rol, V1 trece OK, fără să rescriu logica.'),
        ('6 · AUTONOM', 'STATUS autonom și predarea reală',
         'Autorul scos, zero dependență: STATUS devine AUTONOM singur. Pot lipsi o lună. Pachetul se închide.'),
    ]
    for slide, emotion, phrase in execs:
        h2('SLIDE EXEC ' + slide)
        lbl('STARE (exec-emotion)'); txt(emotion)
        lbl('FRAZĂ DE IMPACT (exec-phrase)'); txt(phrase)

    # ---- DESCHIDERE ----
    h1('DESCHIDERE · DE CE C20')
    txt('Construcția 20 DELEGAREA închide treapta de autonomie și întregul pachet. Construcțiile '
        'dinainte au scos, pe rând, efortul (C18, motorul rulează singur) și supravegherea (C19, '
        'sistemul se ține corect singur). Funcționează. Și totuși, la orice problemă, tot pe autor '
        'îl sună: sistemul e încă proprietatea lui informală. Autonomia e falsă cât timp depinde de '
        'disponibilitatea unei singure persoane.')
    txt('Problema nu este că sistemul nu merge. Problema este că „merge fără mine" nu înseamnă „nu '
        'mai e al meu". Un sistem deținut de o singură persoană moare în ziua în care acea persoană '
        'pleacă. Acum nu mai întrebăm dacă merge, predăm proprietatea unui rol. C20 nu pornește '
        'sistemul și nu îi scrie regulile: ia un sistem guvernat și îi predă proprietatea, dovedind '
        'în workbook că rolul îl poate continua. Motorul vine de mai jos, regulile vin de mai jos; '
        'aici se predă deținerea.')

    # ---- ROLURI ----
    h1('ROLURI · CINE FACE CE')
    lbl('AI (Copilot)')
    txt('Din inventarul de inputuri, separă ce depinde de o singură persoană (author-only, ținut '
        'în cap sau în foi personale) de ce are deja o sursă documentată pe care un rol o poate '
        'folosi. Cu autorul scos, arată exact care input critic încă atârnă de zona autorului. Nu '
        'desemnează nicio persoană, nu scrie reguli de prag, nu preia proprietatea și nu decide '
        'business. AI evidențiază unde mai ești tu cheia; omul mută.')
    lbl('EXCEL · _DELEGARE')
    txt('Găzduiește foaia _DELEGARE: comutatorul AUTOR_ACTIV (semnătura vizuală, scoaterea '
        'autorului), ROL_DELEGAT din listă de roluri, harta de predare (zonă, backup, '
        'responsabilitate, acces pe zone, limite PROTECTED, escaladare), parametrul critic '
        'author-only față de blocul documentat PARAM_, și verificările vii V1-V4 care calculează '
        'STATUS. Totul prin formule care recalculează singure: scoți autorul, lanțul critic '
        'întoarce eroare dacă un input mai atârnă de tine. Valorile sursă rămân neatinse, suma e '
        'conservată cap-coadă.')
    lbl('CONTROL UMAN')
    txt('Inventariem ce ne face indispensabili, definim rolul care preia (rol, nu persoană), '
        'desenăm harta de predare, apăsăm „scoate autorul", citim ce dependență ascunsă se aprinde '
        'și o mutăm în blocul documentat accesibil rolului. Predăm proprietatea și o dovedim în '
        'fișier, nu o explicăm. Nu desemnăm un om, nu scriem o fișă de post, nu facem un curs de '
        'management: predarea trăiește doar pentru că testul e viu.')

    # ---- SCENA REALA ----
    h1('SCENA REALĂ · CE PARE DELEGARE DAR TE LASĂ TOT PROPRIETAR')
    fenomene = [
        ('01 · EXPLICAȚIA',
         'I-ai explicat cum merge, deci crezi că ai predat; dar explicația nu mută proprietatea. '
         'Corecția: predarea se dovedește în workbook prin testul viu, nu printr-un discurs.'),
        ('02 · PROPRIETATE INFORMALĂ',
         'Toți știu că „e al lui", dar nicăieri nu scrie cine îl deține; rămâne o convenție tăcută. '
         'Corecția: proprietatea trece explicit pe un ROL, cu zonă deținută și backup.'),
        ('03 · PERSOANA, NU ROLUL',
         'Predai unui om anume; pleacă omul, pleacă sistemul, fiindcă deținerea s-a legat de o persoană. '
         'Corecția: proprietatea stă pe un rol; persoana se atașează temporar la rol, nu invers.'),
        ('04 · ACCES NECLAR',
         'Cel care ar prelua nu știe ce vede, ce modifică, unde nu are voie; întreabă autorul de fiecare dată. '
         'Corecția: o matrice rol pe zonă spune exact accesul, iar limitele sunt ranges chiar blocate.'),
        ('05 · ESCALADARE VERBALĂ',
         '„Dacă apare ceva, mă suni"; escaladarea trăiește în vorbe, nu se poate dovedi. '
         'Corecția: escaladarea urcă la un ROL, cu un declanșator scris, nu la un semnal de stare (acela e C19).'),
        ('06 · SISTEM ORFAN',
         'Predarea nu poate fi verificată: nimeni nu știe dacă rolul chiar poate opera singur, deci pleci și rămâne orfan. '
         'Corecția: scoți autorul pe viu și sistemul dovedește singur dacă se ține fără tine.'),
    ]
    for titlu, desc in fenomene:
        lbl(titlu); txt(desc)
    txt('Total: 6 fenomene care pretind delegare, dar lasă proprietatea informală la autor, '
        'fiecare cu corecția lui prin testul viu. Recunoaște-le, ca să predai proprietatea unui '
        'rol, nu doar să explici cum merge.')

    # ---- CELE 6 ETAPE ----
    h1('CELE 6 ETAPE · SCRIPT VIDEO')
    etape = [
        {
            'nume': 'ETAPA 1 · Sistemul merge, dar e încă al tău',
            'sub': 'REALITATE',
            'obiectiv': 'Deschidem sistemul moștenit care rulează și se păzește singur, vedem cum la orice '
                        'problemă tot pe autor îl sună, creăm foaia _DELEGARE și o fixăm ca foaia predării.',
            'stare': 'Recunoaștere acută. Sistemul merge fără mine. Și tot eu sunt singurul care îl deține.',
            'voce': '„Motorul rulează, regulile îl țin corect. Și totuși, la orice problemă, telefonul '
                    'sună la mine. Un sistem pe care îl deține o singură persoană nu e autonom, e doar bine '
                    'întreținut de acea persoană."',
            'demo': 'Pe ecran: deschidem Date_MASTER-C20, sistemul guvernat rulează; deschidem _DELEGARE cu '
                    'AUTOR_ACTIV=DA și harta goală, STATUS arată NEPREDAT.',
            'ai': 'Încă nu intervine AI. Doar privim cum sistemul, deși merge, nu are alt deținător decât autorul.',
            'control': 'Punem testul concediului: „Dacă plec o lună, sistemul rămâne orfan? Nu mai întreb dacă '
                       'merge. Predau proprietatea."',
            'reveal': 'Pe ecran: foaia _DELEGARE, fixată ca foaia predării. Aici nu pornim sistemul și nu îi '
                      'scriem regulile, îi predăm proprietatea.',
        },
        {
            'nume': 'ETAPA 2 · Ce te face indispensabil',
            'sub': 'DEPENDENȚĂ',
            'obiectiv': 'Inventariem din _SISTEM ce inputuri atârnă numai de autor, le marcăm author-only, '
                        'cerem Copilot să le separe de cele accesibile rolului, definim rolul care preia. Promptul 1.',
            'stare': 'Privire onestă spre sine. Nu întreb ce face sistemul, întreb ce din el depinde numai de mine.',
            'voce': '„Înainte să predau, văd ce mă face indispensabil: parametri ținuți în cap, foi personale, '
                    'o ajustare lunară pe care doar eu o fac. Pe fiecare pun un flag author-only."',
            'demo': 'Pe ecran: marcăm inputurile AUTHOR_ONLY=DA; deschidem Copilot pe _DELEGARE, lipim Promptul 1; '
                    'AI separă inputurile author-only de cele cu sursă documentată; definim ROL_DELEGAT din lista de roluri.',
            'ai': 'AI parcurge inventarul și împarte inputurile: care depind de o singură persoană și care au o '
                  'sursă pe care un rol o poate folosi. Nu desemnează nicio persoană și nu scrie reguli de prag. AI propune, omul decide.',
            'control': 'Ținem garda rol, nu persoană: „Proprietatea stă pe un rol, de exemplu Operare Raportare '
                       'Vânzări, cu o zonă deținută și un backup-rol. Persoana se atașează temporar la rol. Pleacă omul, rolul rămâne."',
            'reveal': 'Pe ecran: inputurile author-only marcate și ROL_DELEGAT definit din listă de roluri, nu un nume de om.',
        },
        {
            'nume': 'ETAPA 3 · Harta de predare ca input al testului',
            'sub': 'HARTA',
            'obiectiv': 'Scriem responsabilitatea transferată, construim matricea ROL pe ZONĂ pentru acces, '
                        'marcăm limitele PROTECTED, stabilim escaladarea către un rol cu declanșator. Harta alimentează V1-V4.',
            'stare': 'Construcție lucidă. Harta nu e un tabel de citit, e intrarea testului de predare.',
            'voce': '„Scriu ce responsabilitate trece la rol, ce vede și ce modifică pe fiecare zonă, ce nu '
                    'are voie să atingă, și către ce rol urcă o problemă peste mandat. Nu transfer o sarcină, transfer răspunderea."',
            'demo': 'Pe ecran: completăm responsabilitatea, matricea ROLxZONA (acces de operare DA, acces în zone '
                    'protejate NU), marcăm limitele PROTECTED=DA, scriem ESCALADARE_ROL și DECLANSATOR.',
            'ai': 'AI nu intervine aici. Harta e o decizie de proprietate: cine deține ce, până unde, către cine '
                  'urcă. Omul o desenează, sistemul o va testa.',
            'control': 'Ținem garda de scoping: „Escaladarea C20 urcă la un ROL care preia proprietatea, cu un '
                       'declanșator scris, nu la un semnal care schimbă starea. Semnalul de stare era C19. Aici e altă foaie, alt subiect."',
            'reveal': 'Pe ecran: harta de predare completă, fiecare câmp legat de o verificare: accesul alimentează '
                      'V2, limitele alimentează V3, escaladarea alimentează V4.',
        },
        {
            'nume': 'ETAPA 4 · Scoți autorul',
            'sub': 'TESTUL',
            'obiectiv': 'Apăsăm AUTOR_ACTIV=NU, sistemul recalculează singur; cerem Copilot să arate ce input '
                        'critic mai atârnă de autor; confirmăm că STATUS e calculat din V1-V4, nu bifat. Promptul 2.',
            'stare': 'Tensiune controlată. Nu bifez o predare, o dovedesc. Apăs comutatorul și las sistemul să răspundă.',
            'voce': '„Apăs «scoate autorul». Celulele author-only se golesc, formulele care le citeau întorc '
                    'eroare. V1-V4 și STATUS se reașază pe viu. Comutatorul ăsta e semnătura _DELEGARE; C19 nu îl are."',
            'demo': 'Pe ecran: trecem AUTOR_ACTIV pe NU; lipim Promptul 2; AI arată exact ce input critic încă '
                    'atârnă de zona autorului, aprinzând V1; vedem STATUS recalculându-se singur.',
            'ai': 'AI verifică lanțul critic cu autorul scos și evidențiază fiecare dependență rămasă, plus unde '
                  'ar trebui mutat inputul. Nu preia proprietatea și nu decide business. AI evidențiază, omul mută.',
            'control': 'Ținem garda anti-RACI: „STATUS nu e o bifă pe care o pun eu, e o stare calculată din V1-V4. '
                       'Se mișcă singur când schimb harta. O casetă bifată manual ar fi documentare, nu delegare."',
            'reveal': 'Pe ecran: autorul scos, V1-V4 recalculate, STATUS care s-a mutat singur. Testul e viu, nu un tabel.',
        },
        {
            'nume': 'ETAPA 5 · Dependența ascunsă, apoi repararea',
            'sub': 'REPARARE',
            'obiectiv': 'Cu autorul scos, V1 aprinde FAIL pe o dependență ascunsă; mutăm parametrul author-only '
                        'în blocul documentat PARAM_; re-rulăm cu autorul tot scos și nimic nu se mai rupe.',
            'stare': 'Adevărul gol. Credeam că am delegat. De fapt eu eram cheia. Apoi repar, fără să rescriu nimic.',
            'voce': '„Cu autorul scos, o ajustare manuală lunară pe care doar eu o știam lasă o formulă critică '
                    'goală. V1 trece în FAIL, STATUS cade la PARȚIAL. Credeam că am delegat. Tu erai cheia."',
            'demo': 'Pe ecran: la AUTOR_ACTIV=NU și parametrul încă author-only, lanțul critic întoarce eroare, '
                    'V1=FAIL, STATUS=PARȚIAL; mutăm parametrul în blocul documentat PARAM_, accesibil rolului; '
                    'V1 trece în OK, fără să fi schimbat ce calculează sistemul, doar de unde citește.',
            'ai': 'AI a arătat unde era dependența; decizia de a o muta și unde e judecată umană. Repararea nu '
                  'rescrie logica, mută sursa inputului de la autor la rol.',
            'control': 'Ținem granița față de C19: „C19 prindea un INPUT greșit în date, cu reguli și praguri. C20 '
                       'prinde DEPENDENȚA de autor: alt subiect, altă foaie. C19 mă apăra de date proaste; C20 mă scoate pe mine din ecuație."',
            'reveal': 'Pe ecran: parametrul mutat în PARAM_, V1 din FAIL în OK, lanțul critic întreg din surse '
                      'accesibile rolului. Cu autorul tot scos, nimic nu se mai rupe.',
        },
        {
            'nume': 'ETAPA 6 · STATUS autonom și predarea reală',
            'sub': 'AUTONOM',
            'obiectiv': 'Cu autorul scos, acces validat, zone blocate, escaladare către rol și zero dependență '
                        'author-only, STATUS devine AUTONOM singur; predăm proprietatea dovedit în fișier; închidem pachetul.',
            'stare': 'Predare împlinită. Nu explic că rolul poate continua. Sistemul o dovedește singur, fără mine.',
            'voce': '„Apăs «scoate autorul» și sistemul nu se rupe: workbook-ul confirmă singur că alt rol îl '
                    'poate continua. Nu mai sunt cel care explică «uite cum merge». Sunt cel care predă proprietatea, dovedit în fișier."',
            'demo': 'Pe ecran: cu AUTOR_ACTIV=NU și toate V1-V4 OK, STATUS trece în AUTONOM de la sine; testul de '
                    'predare confirmă „autorul e scos și sistemul nu se rupe: alt rol îl poate continua".',
            'ai': 'AI nu mai intervine: dovada e în testul viu, nu într-un discurs. Dacă V1 ar mai aprinde FAIL, '
                  'mai e o dependență de mutat; aici lanțul e întreg.',
            'control': 'Rostim semnătura și granița: „Acum pot lipsi o lună: rolul operează, escaladează, continuă '
                       'fără mine. Sistemul nu mai e al meu, e al rolului care îl deține. C20 închide PACK-ul C01-C20, '
                       'nu mai e treaptă următoare."',
            'reveal': 'Pe ecran: STATUS=AUTONOM cu autorul scos; sistemul predat ca proprietate a unui rol, dovedit '
                      'în testul viu. De la structurarea unui tabel la predarea unui sistem întreg, arcul e complet.',
        },
    ]
    for i, e in enumerate(etape):
        h2(e['nume'])
        lbl('SUBTITLU'); txt(e['sub'])
        lbl('OBIECTIV'); txt(e['obiectiv'])
        lbl('STARE EMOȚIONALĂ'); txt(e['stare'])
        lbl('VOCEA TRAINERULUI'); txt(e['voce'])
        lbl('DEMONSTRAȚIA'); txt(e['demo'])
        lbl('INTERVENȚIA AI'); txt(e['ai'])
        lbl('MOMENT DE CONTROL'); txt(e['control'])
        lbl('REVEAL'); txt(e['reveal'])
        if i < len(etape) - 1:
            lbl('TRANZIȚIE'); txt('Trecem la etapa %d.' % (i + 2))

    # ---- INCHIDERE ----
    h1('ÎNCHIDERE · SISTEMUL CARE NU MAI E AL TĂU')
    txt('C20 predă proprietatea unui sistem de la o persoană la un rol, și o dovedește în workbook. '
        'Foaia _DELEGARE e vie: comutatorul AUTOR_ACTIV scoate autorul, harta de predare alimentează '
        'verificările V1-V4, iar STATUS se calculează singur, NEPREDAT, PARȚIAL, DELEGAT, AUTONOM. '
        'Drama centrală e onestă: cu autorul scos, o dependență ascunsă aprinde V1 în FAIL, dovada că '
        'tu erai cheia; o muți în blocul documentat accesibil rolului și STATUS urcă la AUTONOM, fără '
        'să rescrii logica. C20 nu construiește motorul (vine de jos) și nu scrie regulile (vin de '
        'jos): predă deținerea. Granița față de C19 e clară: C19 prinde un input greșit în date; C20 '
        'prinde dependența de autor. Predăm un sistem deținut de un rol, cu valorile sursă neatinse și '
        'suma conservată cap-coadă. Acum poți lipsi o lună: rolul operează, escaladează, continuă. '
        'C20 închide PACK-ul C01-C20: de la structurarea unui tabel până la predarea unui sistem '
        'întreg, arcul e complet. Nu mai e treaptă următoare.')
    p = doc.add_paragraph()
    r = p.add_run('Pleci, și munca nu mai e a ta.'); r.bold = True; r.font.color.rgb = GOLD
    txt('Trainity · Sistemul 02 Excel · Bogdan Târlă (Dr.Excel)')

    doc.save(OUT)
    print('SCRIS:', OUT)
    print('  paragrafe:', len(doc.paragraphs))


if __name__ == '__main__':
    main()
